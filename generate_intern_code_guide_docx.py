"""
Script to generate the comprehensive, professional .docx documentation:
"Spain Pool Codebase and Architecture: Complete File-by-File & Supervisor Q&A Guide"
Specifically tailored for an intern presenting the entire system to their supervisor.
"""

import os
import json
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PROJECT_ROOT = Path(__file__).resolve().parent
DOCX_OUTPUT_PATH = PROJECT_ROOT / "Spain_Pool_Codebase_and_Architecture_Intern_Guide.docx"

# Color Palette Constants
HEX_NAVY = "0F2C59"       # Primary Headers & Accents
HEX_OCEAN = "3085C3"      # Subheaders & Table Headers
HEX_LIGHT_BLUE = "EAF2F8" # Table Header Alt / Light Background
HEX_CHARCOAL = "1E293B"   # Body Text
HEX_MUTED = "64748B"      # Subtitles & Captions
HEX_BORDER = "CBD5E1"     # Table Borders
HEX_ZEBRA = "F8FAFC"      # Alternating Table Row Shading
HEX_ALERT_RED = "DC2626"  # Red alerts
HEX_ALERT_AMBER = "D97706"# Amber alerts
HEX_ALERT_GREEN = "16A34A"# Green alerts
HEX_CODE_BG = "F1F5F9"    # Code background

COLOR_NAVY = RGBColor(15, 44, 89)
COLOR_OCEAN = RGBColor(48, 133, 195)
COLOR_CHARCOAL = RGBColor(30, 41, 59)
COLOR_MUTED = RGBColor(100, 116, 139)
COLOR_RED = RGBColor(220, 38, 38)
COLOR_AMBER = RGBColor(217, 119, 6)
COLOR_GREEN = RGBColor(22, 163, 74)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_MUTED
    return p


def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_OCEAN
    return p


def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_CHARCOAL
    return p


def add_paragraph(doc, text="", bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_CHARCOAL
    if text:
        r_text = p.add_run(text)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(10)
        r_text.font.italic = italic
        r_text.font.color.rgb = COLOR_CHARCOAL
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_CHARCOAL
    if text:
        r_text = p.add_run(text)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = COLOR_CHARCOAL
    return p


def add_callout(doc, title, text, callout_type="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    
    border_color = HEX_OCEAN
    bg_color = "F0F7FF"
    if callout_type == "WARNING":
        border_color = HEX_ALERT_AMBER
        bg_color = "FFFBEB"
    elif callout_type in ("IMPORTANT", "CAUTION"):
        border_color = HEX_ALERT_RED
        bg_color = "FEF2F2"
    elif callout_type == "SUCCESS":
        border_color = HEX_ALERT_GREEN
        bg_color = "F0FDF4"
    elif callout_type == "INTERN_TIP":
        border_color = "8B5CF6"
        bg_color = "F5F3FF"

    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    prefix = f"[{callout_type}] " if callout_type != "NOTE" else ""
    r_title = p.add_run(f"{prefix}{title}\n")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(10)
    r_title.font.bold = True
    if callout_type in ("IMPORTANT", "CAUTION"):
        r_title.font.color.rgb = COLOR_RED
    elif callout_type == "WARNING":
        r_title.font.color.rgb = COLOR_AMBER
    elif callout_type == "INTERN_TIP":
        r_title.font.color.rgb = RGBColor(139, 92, 246)
    else:
        r_title.font.color.rgb = COLOR_NAVY

    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = COLOR_CHARCOAL

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)


def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, HEX_CODE_BG)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = COLOR_CHARCOAL

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)


def add_styled_table(doc, headers, data, col_widths=None):
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    # Header Row
    hdr_row = tbl.rows[0]
    hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    for idx, heading in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_shading(cell, HEX_NAVY)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(str(heading))
        run.font.name = "Calibri"
        run.font.size = Pt(9.0)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row = tbl.rows[r_idx + 1]
        bg_color = HEX_ZEBRA if (r_idx % 2 == 1) else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=110, right=110)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(8.5)
            run.font.color.rgb = COLOR_CHARCOAL

    if col_widths and len(col_widths) == len(headers):
        for row in tbl.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)
    return tbl


def build_intern_code_guide():
    doc = Document()

    # Set Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # -----------------------------------------------------------------------
    # COVER / HEADER
    # -----------------------------------------------------------------------
    add_title(doc, "POOL PREDICTIVE MAINTENANCE SYSTEM")
    add_subtitle(doc, "Complete Codebase Walkthrough, Database Architecture & Supervisor Q&A Guide\nAuthor: Engineering Intern | Target: Supervisor / Tech Lead Briefing | System Version 6.0")

    add_callout(
        doc,
        "HOW TO USE THIS GUIDE FOR YOUR SUPERVISOR MEETING",
        "This document is designed to give you total technical command over the entire codebase. "
        "It breaks down what every single Python file does, how the PostgreSQL database and Prisma schema work, "
        "how data flows through the system, and includes a dedicated 'Supervisor Q&A Cheat Sheet' with answers to "
        "the toughest technical questions your supervisor is likely to ask.",
        "INTERN_TIP"
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # CHAPTER 1: SYSTEM MENTAL MODEL & DIRECTORY MAP
    # -----------------------------------------------------------------------
    add_heading_1(doc, "1. System Mental Model & Directory Architecture")

    add_paragraph(doc, "The codebase is structured into three clean, decoupled tiers following professional software engineering principles:")
    add_bullet(doc, "Framework-agnostic mathematical modeling, feature engineering, physical kinetics, XGBoost training, and inference chaining. It has zero web or database dependencies, making it 100% testable in isolation.", "1. The ML Core (ml/): ")
    add_bullet(doc, "Production web API built on FastAPI, Prisma ORM, PostgreSQL 16, and APScheduler background jobs. It manages HTTP requests, database persistence, and background tasks.", "2. The Backend Service (backend/): ")
    add_bullet(doc, "React 19 single-page application written in TypeScript and styled with Tailwind CSS, providing operations managers and technicians with interactive fleet monitoring and dosing tools.", "3. The Frontend Dashboard (frontend/): ")

    dir_headers = ["Directory / Module", "Primary Responsibility", "Key Technologies"]
    dir_data = [
        ["ml/config.py", "Single source of truth: regulatory limits, client targets, hyperparameters, setpoints.", "Python dataclass, immutable config"],
        ["ml/features.py", "Pure feature engineering functions (87 signals: lags, headrooms, trends, setpoints).", "Pandas, NumPy, RegEx"],
        ["ml/training/", "Training orchestrator, data cleaning, target interpolation, SHAP, model artifacts.", "XGBoost, Scikit-Learn, SHAP"],
        ["ml/inference/", "Chained multi-day forecaster, physical kinetics decay engine, 525-grid dosing optimizer.", "NumPy, Pickle, JSON"],
        ["backend/store/", "PostgreSQL 16 relational database layer, Prisma schema, migration/seeder scripts, repo queries.", "PostgreSQL 16, Prisma Client Python"],
        ["backend/api/", "FastAPI REST routers: fleet overview, pool details, dosing optimizer, file uploads, admin governance.", "FastAPI, Pydantic V2"],
        ["backend/jobs/", "Async background cron scheduler: daily 4 AM weather refresh, weekly Monday 3 AM retrain.", "APScheduler (AsyncIO)"],
        ["backend/weather/", "Open-Meteo live API consumer with 10-minute in-memory caching and DB fallback.", "Open-Meteo API, Pandas"],
        ["frontend/src/", "Operator dashboard: fleet table, interactive Recharts time series, drag-and-drop file upload modal.", "React 19, TypeScript, Vite, TanStack Query"],
        ["docker/", "Multi-container orchestration: PostgreSQL 16 Alpine, Backend, Migration Seeder, Frontend Nginx.", "Docker, Docker Compose, Nginx"],
        ["tests/", "Comprehensive automated test suite: ML parity, inference bounding, API routes, database operations.", "Pytest, Pytest-AsyncIO, Starlette TestClient"],
    ]
    add_styled_table(doc, dir_headers, dir_data, [Inches(1.8), Inches(3.2), Inches(1.5)])

    # -----------------------------------------------------------------------
    # CHAPTER 2: DATABASE ARCHITECTURE & PRISMA ORM
    # -----------------------------------------------------------------------
    add_heading_1(doc, "2. Database Architecture & Prisma ORM (`backend/store/` & `prisma/`)")

    add_heading_2(doc, "2.1 Why PostgreSQL 16 and Prisma ORM?")
    add_paragraph(doc, "The system uses PostgreSQL 16 Alpine as its primary relational store. To interact with the database in Python without writing fragile raw SQL queries, the system uses Prisma Client Python (`prisma-client-py`). Prisma provides:")
    add_bullet(doc, "Every database query returns strongly-typed Python objects matching the schema.", "Type Safety: ")
    add_bullet(doc, "Prisma's AsyncIO interface matches FastAPI's native async event loop perfectly, avoiding blocking I/O.", "Asynchronous Performance: ")
    add_bullet(doc, "Changes to `prisma/schema.prisma` are pushed into PostgreSQL via automated migrations (`prisma db push`).", "Declarative Migrations: ")

    add_heading_2(doc, "2.2 Complete Database Schema Models (`prisma/schema.prisma`)")
    
    schema_headers = ["Model (Table)", "Primary Key & Indexes", "Fields & Data Types", "Business Purpose"]
    schema_data = [
        [
            "Pool\n(pools)",
            "PK: pool_id (String)\nIndex: pool_id",
            "community_name (String?)\npool_type, deck_type (String?)\npool_volume_m3, pool_surface_m2 (Float?)\nfilter_diameter, filter_count, motor_count (Float?)\npool flags: heated, community, outdoor, skimmer, overflow (Int 0/1)\ndeck flags: deck_grass, deck_mixed, deck_paved (Float)",
            "Stores static physical attributes and equipment specifications for all registered community pools. One Pool has many Readings."
        ],
        [
            "Reading\n(readings)",
            "PK: id (Int autoincrement)\nUnique: [pool_id, reading_date]\nIndex: [pool_id, reading_date]",
            "pool_id (String FK -> Pool cascade)\nreading_date (DateTime)\ntechnician (String?)\nph, free_chlorine, turbidity (Float?)\nhypochlorite_dosing_pct, hypochlorite_dosing_hours (Float?)\nph_dosing_pct, ph_dosing_hours (Float?)\ndaily_filtration_hours, water_temperature (Float?)\nsource (String default 'upload')\ncreated_at (DateTime now)",
            "Stores historical and newly ingested water quality measurements and pump settings. Unique constraint prevents duplicate readings for the same pool on the same date."
        ],
        [
            "WeatherDaily\n(weather_daily)",
            "PK: date (DateTime)",
            "w_temp_max, w_temp_mean (Float?)\nw_uv_max, w_uv_clear_sky_max (Float?)\nw_solar_radiation, w_sunshine_hours (Float?)\nw_precipitation_mm, w_wind_max_kmh (Float?)\nw_et0 (Float?)\nw_weather_code (Int?)\nfetched_at (DateTime now)",
            "Daily atmospheric record for Alicante (Lat 38.3452, Lon -0.4815). Used for solar photolysis, temperature degassing, and forward weather forecasts."
        ],
        [
            "ModelRun\n(model_runs)",
            "PK: run_id (String)",
            "artifact_dir (String)\ncreated_at (DateTime now)\nis_active (Int 0 or 1)\nmetrics_json (String?)\nfeature_schema_json (String?)\npromoted_at (DateTime?)\npromote_reason (String?)",
            "Model registry tracking all historical training runs, holdout test metrics (RMSE, MAE, R²), and which run is actively serving live predictions."
        ],
        [
            "IngestLog\n(ingest_logs)",
            "PK: id (Int autoincrement)",
            "source (String: upload, manual, ingest_api, master)\nfilename (String?)\npool_count, row_count, skipped_count (Int)\ncreated_at (DateTime now)\ndetail_json (String? storing skipped row diagnostics)",
            "Immutable audit trail logging every data ingestion event, recording how many rows were inserted and details of any malformed rows skipped."
        ],
    ]
    add_styled_table(doc, schema_headers, schema_data, [Inches(1.2), Inches(1.5), Inches(2.2), Inches(1.6)])

    add_heading_2(doc, "2.3 Database Python Files Walkthrough")
    
    add_paragraph(doc, "1. `backend/store/client.py`: Manages the database connection lifecycle. It creates a global Prisma client instance, exposes `connect_db()` and `disconnect_db()` for FastAPI's lifespan, and provides `get_db()` as a FastAPI dependency.")
    add_paragraph(doc, "2. `backend/store/schema.py`: Contains pure Pydantic and data-transfer representations used across repository boundaries.")
    add_paragraph(doc, "3. `backend/store/repo.py`: The database repository layer containing all optimized queries:")
    add_bullet(doc, "Retrieves the single latest reading for a pool joined with its physical pool metadata, constructing the base Series for prediction.", "`get_master_row(pool_id)`: ")
    add_bullet(doc, "Returns pool IDs with recorded readings in the past 30 days.", "`get_active_pool_ids(as_of_date)`: ")
    add_bullet(doc, "Performs atomic bulk inserts/updates of readings with conflict handling.", "`upsert_readings_batch(rows, source)`: ")
    add_bullet(doc, "Synchronizes daily weather records into `weather_daily`.", "`upsert_weather_batch(rows)`: ")
    add_bullet(doc, "Atomically switches the active model pointer and records promotion reasoning.", "`set_active_model_run(run_id, reason)`: ")
    add_paragraph(doc, "4. `backend/store/migrate.py`: Standalone migration and seeder script. It runs `prisma db push` via subprocess to ensure PostgreSQL matches the schema, reads `outputs/master_dataset_v6.csv`, deduplicates entries, and seeds pools, historical readings, weather records, and active model runs.")

    # -----------------------------------------------------------------------
    # CHAPTER 3: MACHINE LEARNING PACKAGE WALKTHROUGH
    # -----------------------------------------------------------------------
    add_heading_1(doc, "3. Machine Learning Package Deep Dive (`ml/`)")

    add_paragraph(doc, "The `ml/` package contains the predictive brain of the system. It is divided into two distinct sub-packages: `ml/training/` (for building models) and `ml/inference/` (for running live forecasts in production).")

    add_heading_2(doc, "3.1 `ml/config.py` — The Single Source of Truth")
    add_paragraph(doc, "This file defines all immutable configuration constants and the `PipelineConfig` dataclass:")
    add_bullet(doc, "Real Decreto 742/2013 thresholds (Cl min: 0.5, Cl close: 5.0, pH: 7.2–8.0, Turbidity: 5.0) and client optimal targets (Cl: 1.0–1.5, pH: 7.4).", "Regulatory & Target Constants: ")
    add_bullet(doc, "The assumed post-treatment ideal (Cl: 2.5, pH: 7.4, Turb: 0.5) from which inter-visit degradation evolves.", "Post-Treatment Setpoints: ")
    add_bullet(doc, "56-key dictionary mapping messy Spanish spreadsheet headers to clean snake_case names.", "RENAME_MAP: ")
    add_bullet(doc, "Holds filepaths, Alicante GPS coordinates (38.3452° N, -0.4815° W), XGBoost hyperparameters, and promotion gate tolerances.", "PipelineConfig Dataclass: ")

    add_heading_2(doc, "3.2 `ml/features.py` — Feature Engineering Engine")
    add_paragraph(doc, "Contains all feature calculation functions. Because these functions are imported by both training and inference, feature drift is mathematically impossible:")
    add_bullet(doc, "Computes linear distances from current readings to the 0.5 pathogen floor, 5.0 closure ceiling, and 7.2/8.0 pH boundaries.", "`add_headroom_features(df)`: ")
    add_bullet(doc, "Computes directional changes (today - lag1) and daily velocity (trend / days_since_last_visit).", "`add_trend_features(df)`: ")
    add_bullet(doc, "Computes degradation deltas (setpoint - current) and per-day degradation rates vs inter-visit gap.", "`add_setpoint_features(df)`: ")
    add_bullet(doc, "Calculates active hypochlorous acid (HOCl) fraction as a function of pH (chlorine loses disinfection efficacy as pH exceeds 7.5).", "`cl_effectiveness(cl, ph)`: ")

    add_heading_2(doc, "3.3 `ml/training/` — Model Training Pipeline")
    add_paragraph(doc, "1. `ml/training/steps.py`: Contains pure pandas-in / pandas-out functions for each stage of data preparation:")
    add_bullet(doc, "Loads master Excel and renames Spanish headers.", "STEP 1 (load_and_rename): ")
    add_bullet(doc, "Filters dataset strictly to the 135 pools equipped with liquid chlorine dosing pumps.", "STEP 1.5 (filter_chlorine_pump_pools): ")
    add_bullet(doc, "Loads or downloads Alicante daily weather.", "STEP 2 (load_or_fetch_weather): ")
    add_bullet(doc, "Separates side-by-side readings, ops, and product tables.", "STEP 3 (separate_subtables): ")
    add_bullet(doc, "Cleans readings, operations, and products; deduplicates multi-visit days.", "STEP 4 (clean_*): ")
    add_bullet(doc, "Backfills physical pool dimensions using per-pool max and fleet medians.", "STEP 4.5 (backfill_static): ")
    add_bullet(doc, "As-of backward merge of ops and chemical products with 14-day tolerance.", "STEP 5 (merge_subtables): ")
    add_bullet(doc, "Joins today's weather and tomorrow's forecast (shifted by -1 day).", "STEP 6 (join_weather): ")
    add_bullet(doc, "Computes all 87 features and cumulative weather since last visit.", "STEP 7 (engineer_features): ")
    add_bullet(doc, "Formulates synthetic next-day targets via linear 1/k setpoint interpolation.", "STEP 8 (build_targets): ")
    add_bullet(doc, "Executes 80/20 chronological split (cutoff: Oct 13, 2025) and fits Scikit-Learn OneHotEncoder/ColumnTransformer preprocessor.", "STEP 9 (select_features_and_split): ")

    add_paragraph(doc, "2. `ml/training/train.py`: The pipeline orchestrator. Fits three XGBoost regressors (Model A: Free Chlorine, Model C: pH, Model D: Turbidity), runs SHAP TreeExplainer, writes artifacts to `models/<run_id>/`, evaluates candidate metrics against the active run via `ml/training/evaluate.py`, and updates `models/latest.json` if promoted.")

    add_paragraph(doc, "3. `ml/training/artifacts.py`: Implements `ArtifactStore`. Writes models (`xgb_*.json`), preprocessor (`preprocessor_v6.pkl`), and metadata (`inference_config_v6.json`) into an atomic `.tmp` directory before renaming, preventing corrupt partial writes.")

    add_heading_2(doc, "3.4 `ml/inference/` — Operational Inference & Dosing")
    add_paragraph(doc, "1. `ml/inference/predictor.py`: Implements the `predict_forward` pure function and `PredictionService` class. Loads the active model run once in memory, provides graceful fallback on errors, and supports dynamic hot-reloading when retrain jobs finish.")
    add_paragraph(doc, "2. `ml/inference/chaining.py`: Implements the chained multi-day forecast engine. Rolls pool states forward day-by-day, recalculates autoregressive features, blends predictions with physical kinetics equations, computes uncertainty bands (± MAE * sqrt(step)), and assigns operational urgency (🚨 Immediate, ⚠️ Advised, ⚠️ Monitor, ✅ Routine).")
    add_paragraph(doc, "3. `ml/inference/optimiser.py`: Implements the `Optimiser` class. Evaluates 525 candidate dosing pump configurations (dosing % 0–100% × hours 0–24h) using `_FeatureEnv` to find the minimal chemical effort that keeps water quality within client target ranges.")

    # -----------------------------------------------------------------------
    # CHAPTER 4: BACKEND & API WALKTHROUGH
    # -----------------------------------------------------------------------
    add_heading_1(doc, "4. Backend Architecture & REST API Walkthrough (`backend/`)")

    add_heading_2(doc, "4.1 Application Lifecycle (`backend/main.py`)")
    add_paragraph(doc, "The FastAPI entrypoint coordinates application startup and shutdown via an async `lifespan` manager:")
    add_code_block(doc,
        "@asynccontextmanager\n"
        "async def lifespan(app: FastAPI):\n"
        "    # 1. Connect to PostgreSQL via Prisma\n"
        "    await connect_db()\n"
        "    # 2. Warm in-memory weather cache from database\n"
        "    await warm_weather_cache(client=db)\n"
        "    # 3. Load active ML model artifacts into PredictionService\n"
        "    svc = PredictionService(settings.models_dir_path)\n"
        "    svc.load()\n"
        "    app.state.prediction_service = svc\n"
        "    # 4. Start background APScheduler (weather + retrain jobs)\n"
        "    if settings.enable_scheduler:\n"
        "        start_scheduler(settings)\n"
        "    yield\n"
        "    # Shutdown: Stop scheduler and disconnect DB cleanly\n"
        "    if settings.enable_scheduler:\n"
        "        shutdown_scheduler()\n"
        "    await disconnect_db()"
    )

    add_paragraph(doc, "Middleware & Tracing: `request_tracing_middleware` assigns a unique UUID `X-Request-ID` to every HTTP request, measures execution duration, attaches `X-Response-Time`, and logs method, path, status, and duration.")

    add_heading_2(doc, "4.2 API Routers (`backend/api/`)")
    add_paragraph(doc, "1. `backend/api/fleet.py`: `GET /api/fleet` endpoint. Aggregates chained forecasts across all active pools for an `as_of` date, extracts Today and Tomorrow forecast chips, computes urgency, and returns paginated results.")
    add_paragraph(doc, "2. `backend/api/pool.py`: `GET /api/pool/{pool_id}` endpoint. Delivers complete pool analytics: latest measurements, chained multi-day forecast with uncertainty bands, 525-grid dosing recommendation, and historical time-series points for Recharts.")
    add_paragraph(doc, "3. `backend/api/optimise.py`: `GET /api/optimise/{pool_id}` endpoint. Standalone dosing recommendation endpoint.")
    add_paragraph(doc, "4. `backend/api/upload.py`: Handles file uploads (`POST /api/upload`) and manual reading creation (`POST /api/readings`). Uses `_auto_detect_mapping` with `SequenceMatcher` to automatically identify column names from uploaded Excel or CSV files, caches parsed frames to temporary files, and executes batch upserts upon user confirmation (`POST /api/map-columns`).")
    add_paragraph(doc, "5. `backend/api/ingest.py`: Programmatic ingestion endpoints (`POST /api/ingest/readings` and `POST /api/ingest/readings/file`) for external sensor feeds, telemetry gateways, or automated scripts, protected by Bearer token authorization.")
    add_paragraph(doc, "6. `backend/api/admin.py`: Protected administrative endpoints (`/api/admin/runs`, `/api/admin/retrain`, `/api/admin/weather-refresh`, `/api/admin/ingest-log`) for triggering manual retraining, weather synchronization, and reviewing audit logs.")
    add_paragraph(doc, "7. `backend/api/health.py`: Health probes (`/healthz`, `/healthz/live`, `/healthz/ready`) returning HTTP 200/503 for Kubernetes and Docker container monitoring.")

    add_heading_2(doc, "4.3 Background Jobs & Weather Provider")
    add_paragraph(doc, "1. `backend/jobs/scheduler.py`: Configures APScheduler `AsyncIOScheduler`. Runs weather refresh daily at 4:00 AM (`0 4 * * *`) and periodic retrain weekly on Monday at 3:00 AM (`0 3 * * 1`).")
    add_paragraph(doc, "2. `backend/jobs/retrain.py`: Retraining worker. Checks if >= 200 new readings have accumulated, executes `ml.training.train` as a non-blocking subprocess with timeout protection, evaluates the candidate model against the promotion gate, registers the run in PostgreSQL `model_runs`, and hot-reloads the active model.")
    add_paragraph(doc, "3. `backend/weather/provider.py`: Live weather provider. Fetches yesterday's archive and 7-day forecast from Open-Meteo, upserts into `weather_daily`, and maintains a fast in-memory dictionary cache with a 10-minute TTL to power high-speed inference lookups.")

    # -----------------------------------------------------------------------
    # CHAPTER 5: FRONTEND ARCHITECTURE
    # -----------------------------------------------------------------------
    add_heading_1(doc, "5. Frontend Dashboard Architecture (`frontend/`)")

    add_paragraph(doc, "The user interface is built with React 19, TypeScript, and Tailwind CSS, bundled with Vite:")
    add_bullet(doc, "Central fleet view displaying KPI summary cards (Immediate, Advised, Routine, Extended), search/filter toolbar, and the interactive fleet table with color-coded Today/Tomorrow forecast chips.", "`frontend/src/pages/FleetPage.tsx`: ")
    add_bullet(doc, "Deep-dive pool analytics view featuring interactive horizon selectors (2, 3, 5, 7 days), chemical dosing pump recommendation cards, and three Recharts line graphs with shaded regulatory safety bands.", "`frontend/src/pages/PoolDetailPage.tsx`: ")
    add_bullet(doc, "Model lifecycle control panel displaying the active model run, test metrics, manual retrain/weather buttons with live animated spinners, model run registry table, and data ingestion audit log.", "`frontend/src/pages/AdminPage.tsx`: ")
    add_bullet(doc, "Modal supporting drag-and-drop Excel/CSV file upload with automatic column mapping, manual reading entry form, and REST API documentation.", "`frontend/src/components/IngestModal.tsx`: ")
    add_bullet(doc, "Type-safe async fetch wrapper exposing clean TypeScript functions (`api.fleet()`, `api.pool()`, `api.upload()`, etc.) consumed by TanStack React Query hooks.", "`frontend/src/api.ts`: ")

    # -----------------------------------------------------------------------
    # CHAPTER 6: STEP-BY-STEP DATA FLOW SCENARIOS
    # -----------------------------------------------------------------------
    add_heading_1(doc, "6. End-to-End Execution Scenarios (What Happens Under the Hood)")

    add_heading_2(doc, "6.1 Scenario A: An Operator Uploads an Excel File")
    add_bullet(doc, "User selects an Excel file (.xlsx) in the `IngestModal` component.", "1. File Selection: ")
    add_bullet(doc, "`api.uploadFile()` sends multipart file to `POST /api/upload`.", "2. Upload Request: ")
    add_bullet(doc, "`upload.py` parses Excel with openpyxl, extracts columns, runs `_auto_detect_mapping` via string similarity, caches the parsed table to a temporary file (`upload_id.json`), and returns a 5-row preview with suggested column mappings.", "3. Parsing & Mapping Detection: ")
    add_bullet(doc, "Operator confirms or adjusts column selections and clicks 'Import Readings'.", "4. User Confirmation: ")
    add_bullet(doc, "`POST /api/map-columns` reads cached file, parses dates, validates numeric measurements, upserts pools into `pools`, executes bulk batch upsert into `readings`, records event in `ingest_logs`, deletes temp file, and returns success count.", "5. Ingestion & Persistence: ")

    add_heading_2(doc, "6.2 Scenario B: A User Opens the Fleet Page")
    add_bullet(doc, "`FleetPage.tsx` fires TanStack `useQuery` targeting `GET /api/fleet?page=0&page_size=50`.", "1. UI Request: ")
    add_bullet(doc, "`fleet.py` queries `repo.get_active_pool_ids()` to find pools visited in the last 30 days.", "2. Pool Discovery: ")
    add_bullet(doc, "For each pool, `repo.get_master_row()` fetches latest reading and physical attributes.", "3. Master Row Extraction: ")
    add_bullet(doc, "`PredictionService.forecast()` executes `predict_forward`, performing chained day-by-day rollout from last visit date to tomorrow, injecting cached weather from `_weather_cache`.", "4. Chained Forecast Rollout: ")
    add_bullet(doc, "Extracts Today/Tomorrow states, evaluates regulatory breaches, sorts by urgency (Immediate -> Routine), and returns paginated JSON.", "5. Urgency Sorting & Response: ")

    add_heading_2(doc, "6.3 Scenario C: Scheduled Retraining on Monday at 3:00 AM")
    add_bullet(doc, "APScheduler fires `_retrain_job` in `scheduler.py`.", "1. Trigger: ")
    add_bullet(doc, "`retrain.py` checks `should_retrain()`. If new readings >= 200, it acquires `_retrain_lock`.", "2. Threshold Gate: ")
    add_bullet(doc, "Spawns `python -m ml.training.train --run-id v6-YYYYMMDD-HHMMSS` as an async subprocess.", "3. Subprocess Execution: ")
    add_bullet(doc, "`train.py` executes Steps 1–12, trains 3 XGBoost models, computes test RMSE/MAE, and writes artifacts into `models/<run_id>/`.", "4. Model Fitting: ")
    add_bullet(doc, "`evaluate.py` evaluates candidate vs active metrics. If holdout MAE is within promotion slack, it updates `models/latest.json` and PostgreSQL `model_runs`.", "5. Promotion & Registry: ")
    add_bullet(doc, "`PredictionService.reload()` hot-reloads the new model into memory with zero downtime.", "6. Hot-Reload: ")


    # -----------------------------------------------------------------------
    # CHAPTER 7: SUPERVISOR Q&A CHEAT SHEET
    # -----------------------------------------------------------------------
    add_heading_1(doc, "7. Supervisor Q&A Preparation Guide (Intern Cheat Sheet)")

    add_paragraph(doc, "Use these detailed answers when your supervisor or tech lead asks deep technical questions during review:")

    qa_list = [
        (
            "Q1: Why did we choose XGBoost instead of a Deep Learning model (like LSTM or Transformer)?",
            "1. Tabular Domain Dominance: Tree-based gradient boosting consistently outperforms neural networks on tabular datasets with heterogeneous features (continuous chemistry, discrete calendar flags, categorical pool types).\n"
            "2. Small Sample Efficiency: With 135 pools and ~38,000 readings, deep learning models are prone to severe overfitting, whereas XGBoost regularizes exceptionally well via max_depth=5, reg_alpha=0.1, and reg_lambda=1.0.\n"
            "3. Explainability: XGBoost integrates natively with SHAP TreeExplainer, allowing us to explain exact chemical feature drivers for every prediction to client operators."
        ),
        (
            "Q2: How does the system handle irregular technician visit gaps (e.g. 2 days vs. 5 days)?",
            "We handle irregular visit gaps in two places:\n"
            "1. In Training (Target Formulation): Targets are scaled by (1 / k) via linear interpolation: Target = Setpoint + (Next_Reading - Setpoint) / k. This normalizes every training target to represent exactly 1 day of degradation, regardless of how many days elapsed.\n"
            "2. In Production Inference (Chained Forecaster): The chained predictor steps forward day-by-day from the last visit date, recalculating autoregressive lags and injecting daily weather at each step until it reaches Today and Tomorrow."
        ),
        (
            "Q3: Why did the R² score drop in V6 compared to earlier versions (e.g. V5)?",
            "In earlier versions, synthetic targets were defined as (Reading_Today - Decay), which made the target almost a copy of the input reading (inflating R² to ~0.80 falsely due to high autocorrelation). In V6, we re-anchored degradation to the post-treatment setpoint (Cl 2.5, pH 7.4), reflecting the real-world 'measure -> treat -> degrade' cycle. The new targets have genuine variation, which honestly lowered R² (~0.26) while improving actual Mean Absolute Error (MAE improved from 0.26 to 0.1972 mg/L)."
        ),
        (
            "Q4: How do we prevent temporal data leakage during training?",
            "We strictly avoid random K-fold cross-validation. Instead, we enforce a strict 80/20 chronological split at the 80th percentile timestamp (October 13, 2025). The model trains only on historical data (2023 to Oct 2025) and is evaluated exclusively on future unseen data (Oct 2025 to Aug 2026)."
        ),
        (
            "Q5: How does the system prevent non-physical predictions (e.g. chlorine dropping below 0 or rising without chemicals)?",
            "We implement a Hybrid Physical Kinetics Rate Integration Engine. For example, chlorine decay is bounded by first-order solar photolysis: Cl_kinetic = Anchor_Cl * exp(-decay_k / 3.0), where decay_k scales with solar radiation. Unless chemicals were added, the prediction is bounded by this physical kinetic ceiling. Similarly, pH upward drift is bounded by temperature-driven CO2 degassing."
        ),
        (
            "Q6: How does the backend update models without restarting or dropping requests?",
            "PredictionService loads model artifacts into memory and holds thread-safe references. When a retraining job completes and passes the promotion gate, it calls PredictionService.reload(), which loads the new models from models/latest.json into new variables before swapping references atomically. If a reload fails, it logs an error and keeps the existing model in memory, ensuring zero downtime."
        ),
        (
            "Q7: What happens if the Open-Meteo weather API fails or is unreachable?",
            "1. In-Memory Cache: The weather provider checks its in-memory cache first (10-minute TTL).\n"
            "2. Database Fallback: If not in memory, it queries the local PostgreSQL weather_daily table (which holds 1,312 historical and forecasted days).\n"
            "3. Imputation Fallback: If a specific weather date is completely missing, the feature preprocessor automatically fills missing values with training medians (e.g. solar radiation: 22.69 MJ/m², temperature: 24.0°C), so predictions never crash."
        ),
        (
            "Q8: How does the database prevent duplicate readings for the same pool on the same day?",
            "In prisma/schema.prisma, the Reading model defines a compound unique constraint: @@unique([pool_id, reading_date]). When data is ingested via upsert_readings_batch, Prisma executes an UPSERT (INSERT ... ON CONFLICT DO UPDATE), updating existing records rather than creating duplicates."
        ),
        (
            "Q9: How does the Dosing Optimizer balance chemical dosage percentage vs pump run hours?",
            "It evaluates 525 discrete combinations (hypochlorite % 0–100% in 5% steps × hours 0–24h in 1h steps). It calculates a penalty based on distance outside the client target band [Cl 1.0–1.5, pH 7.2–8.0]. For all feasible configurations (penalty = 0), it uses chemical effort cost = (dosing_pct / 100) * hours as the tie-breaker, recommending the lowest chemical and pump wear option."
        ),
        (
            "Q10: How do you guarantee feature parity between training and production inference?",
            "All feature extraction math is implemented in pure, reusable functions in ml/features.py (e.g. add_headroom_features, add_trend_features, add_setpoint_features). Both ml/training/steps.py and ml/inference/predictor.py import the exact same functions. Furthermore, tests/ml/test_feature_parity.py runs golden tests verifying bit-for-bit parity."
        ),
    ]

    for q, a in qa_list:
        add_heading_2(doc, q)
        add_paragraph(doc, a)

    # -----------------------------------------------------------------------
    # CHAPTER 8: QUICK COMMANDS REFERENCE
    # -----------------------------------------------------------------------
    add_heading_1(doc, "8. Quick Developer & CLI Commands Reference")

    cmd_headers = ["Action", "Terminal Command", "What It Does"]
    cmd_data = [
        ["Run Test Suite", "pytest -v", "Runs all unit and integration tests across API, ML, and Store modules."],
        ["Run Training Pipeline", "python -m ml.training.train", "Executes full training pipeline, evaluates promotion gate, and updates latest model pointer."],
        ["Dry-Run Pipeline Test", "python -m ml.training.train --dry-run", "Validates dataset loading and pool filtering without full model training."],
        ["Run Inference CLI", "python inference.py", "Executes chained daily forecast across all active pools and prints formatted console table."],
        ["Start Backend Dev Server", "uvicorn backend.main:app --reload --port 8000", "Starts FastAPI ASGI development server with live auto-reload."],
        ["Start Frontend Dev Server", "cd frontend && npm run dev", "Starts React Vite development server on http://localhost:5173."],
        ["Run Database Migration/Seed", "python -m backend.store.migrate", "Synchronizes Prisma schema with PostgreSQL and seeds pools, readings, and weather."],
        ["Start Full Docker Stack", "docker-compose up --build", "Spawns PostgreSQL 16, migration seeder, FastAPI backend, and Nginx frontend."],
    ]
    add_styled_table(doc, cmd_headers, cmd_data, [Inches(1.8), Inches(2.7), Inches(2.0)])

    # Save document
    doc.save(str(DOCX_OUTPUT_PATH))
    print(f"Intern Guide successfully generated at: {DOCX_OUTPUT_PATH}")


if __name__ == "__main__":
    build_intern_code_guide()
