#!/usr/bin/env python3
"""
Generate the Data Column Feasibility Report (DOCX) for the
Smart Predictive Pool Maintenance System.

Maps every raw column in merged_pool_data_2017_2022.csv to:
  - English translation
  - Workflow document prediction relevance
  - Data availability
  - Whether the prediction can be built

Output: data_feasibility_report.docx
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ─────────────────────────────────────────────────
# Load and analyse the dataset
# ─────────────────────────────────────────────────
CSV_PATH = "data/merged_pool_data_2017_2022.csv"
df = pd.read_csv(CSV_PATH)
total_rows = len(df)
total_pools = df["PISCINA"].nunique()
dates = pd.to_datetime(df["FECHA"], format="mixed", dayfirst=True, errors="coerce")
date_min = dates.min()
date_max = dates.max()

def pool_count(col):
    """Number of unique pools that have at least one non-null value."""
    if col not in df.columns:
        return 0
    return df[df[col].notna()]["PISCINA"].nunique()

def non_null(col):
    if col not in df.columns:
        return 0
    return int(df[col].notna().sum())

def fill_pct(col):
    nn = non_null(col)
    return round(nn / total_rows * 100, 1) if total_rows > 0 else 0.0

def per_pool_stats(col):
    """Returns (min, median, max, pools_gte_10) for non-null values per pool."""
    if col not in df.columns or non_null(col) == 0:
        return (0, 0, 0, 0)
    per = df[df[col].notna()].groupby("PISCINA").size()
    return (int(per.min()), int(per.median()), int(per.max()), int((per >= 10).sum()))


# ─────────────────────────────────────────────────
# Column definitions — THE BIG TABLE
# ─────────────────────────────────────────────────
# Each entry: (spanish_name, english_name, field_type, relevance, predictions_affected, can_build_note)
# field_type: "core_reading", "static_pool", "static_equipment", "dynamic_ops", "product", "meta", "separator"

columns_info = [
    # ── SUB-TABLE 1: Water Quality Readings ──
    ("PISCINA", "Pool ID", "meta",
     "Primary key — identifies each pool uniquely. Essential for all per-pool grouping, lag computation, and model training.",
     "All predictions (P1–P5)",
     "✅ YES — 96.7% filled, 522 unique pools"),

    ("COMUNIDAD", "Community / Address", "meta",
     "Identifies the community or building the pool belongs to. Not used as an ML feature but useful for operational reporting and route planning.",
     "None (operational context only)",
     "N/A — Not an ML feature"),

    ("FECHA", "Reading Date", "core_reading",
     "HIGH — The timestamp of each water quality reading. Critical for computing all temporal features: days_since_last_visit, visit_month, is_summer, and all rate-per-day calculations. Also used for the temporal train/test split.",
     "All predictions (P1–P5)",
     "✅ YES — 92.4% filled"),

    ("EMPLEADO", "Technician Name", "meta",
     "LOW — Records which technician took the reading. Could be used to model technician-specific measurement bias, but not currently used as a feature.",
     "None currently (potential future feature)",
     "✅ Available — 88.3% filled"),

    ("PH", "pH Level", "core_reading",
     "CRITICAL — Measures water acidity/alkalinity. Primary target for pH prediction model. Used to derive: pH lags, pH rolling mean/std, pH headroom (distance from 7.2–8.0 limits), pH trend, pH drift rate, and pH-Chlorine Effectiveness Index.",
     "P1 (Chlorine Safety — pH affects chlorine effectiveness), P4 (Chemical Cost — pH drift drives corrector demand)",
     "✅ YES — 91.6% filled, 469 pools with ≥20 readings"),

    ("TURBIDEZ", "Turbidity (NTU)", "core_reading",
     "CRITICAL — Measures water cloudiness. Primary target for turbidity prediction model. Used to derive: turbidity lags, rolling mean, turbidity headroom (distance from 5.0 NTU limit), turbidity growth rate.",
     "P2 (Filter Backwash — turbidity rise indicates filter loading)",
     "✅ YES — 70.7% filled, 460 pools with ≥20 readings"),

    ("CLORO LIBRE", "Free Chlorine (mg/L)", "core_reading",
     "CRITICAL — Measures disinfectant level. Primary target for chlorine prediction model. Used to derive: chlorine lags, rolling mean/std, chlorine headroom (distance from 0.5 and 5.0 limits), chlorine decay rate, chlorine deficit.",
     "P1 (Chlorine Safety — direct input), P4 (Chemical Cost — chlorine demand forecasting)",
     "✅ YES — 90.1% filled, 468 pools with ≥20 readings"),

    ("Unnamed: 7", "— Blank separator —", "separator",
     "N/A — Empty column separating Sub-table 1 from Sub-table 2 in the original spreadsheet.",
     "None", "N/A"),

    ("ABUSO CREMAS PROTECCION", "Sunscreen Contamination Flag", "static_pool",
     "LOW — Flags pools where excessive sunscreen use by bathers is a known issue. Sunscreen creates organic load that consumes chlorine faster.",
     "P1 (Chlorine Safety — explains abnormal chlorine consumption)",
     "⚠️ PARTIAL — Only 36/522 pools flagged (0.1% of rows). Too sparse for reliable ML signal."),

    ("Caudal bomba de PH", "pH Pump Flow Rate (L/h)", "static_equipment",
     "HIGH (for equipment prediction) — The designed/actual flow rate of the pH dosing pump. Used as baseline to compute pH Pump Hydraulic Efficiency = Actual ÷ Designed × 100.",
     "P3 (Pump Degradation), P5 (pH Dosing Pump Maintenance Cycle)",
     "❌ NO — Only 68/522 pools (0.2% of rows). Static field, but covers only 13% of the fleet."),

    ("Caudal bomba hipoclorito", "Chlorine Pump Flow Rate (L/h)", "static_equipment",
     "HIGH (for equipment prediction) — The designed/actual flow rate of the chlorine dosing pump. Used as baseline to compute Chlorine Pump Hydraulic Efficiency.",
     "P3 (Pump Degradation), P5 (Chlorine Dosing Pump Maintenance Cycle)",
     "⚠️ PARTIAL — 159/522 pools (0.5% of rows). Covers 30% of the fleet."),

    ("Caudal del motor", "Main Pump Flow Rate (m³/h)", "static_equipment",
     "HIGH (for equipment prediction) — The designed/actual flow rate of the main circulation pump. Used to compute Hydraulic Efficiency and Turnover Time = Pool Volume ÷ Pump Flow Rate.",
     "P2 (Filter Backwash — flow + turbidity), P3 (Pump Degradation), P5 (Main Pump Maintenance Cycle)",
     "⚠️ PARTIAL — 254/522 pools (0.7% of rows). Covers ~49% of the fleet. Static, so backfillable."),

    ("Diametro filtro", "Filter Diameter (mm)", "static_equipment",
     "MEDIUM — Defines physical filtration capacity. Used to set per-pool thresholds for filter maintenance and replacement alerts.",
     "P2 (Filter Backwash — defines capacity thresholds), P5 (Filter Media Maintenance Cycle)",
     "⚠️ PARTIAL — 313/522 pools (0.9% of rows). Covers ~60% of the fleet. Static, backfillable."),

    ("Numero de filtros", "Number of Filters", "static_equipment",
     "MEDIUM — Pools with multiple filters have more capacity and different degradation thresholds.",
     "P2 (Filter Backwash), P5 (Filter Media Maintenance Cycle)",
     "⚠️ PARTIAL — 267/522 pools (0.8% of rows). Covers ~51% of the fleet. Static, backfillable."),

    ("Número de motores", "Number of Motors/Pumps", "static_equipment",
     "MEDIUM — Determines alert priority: single-pump pools have no backup, so degradation is more urgent.",
     "P3 (Pump Degradation — urgency level), P5 (Main Pump Maintenance Cycle — priority escalation)",
     "⚠️ PARTIAL — 194/522 pools (0.6% of rows). Covers ~37% of the fleet. Static, backfillable."),

    ("PISCINA CLIMATIZADA", "Heated Pool (flag)", "static_pool",
     "LOW — Heated pools have different chlorine decay and temperature profiles.",
     "P1 (Chlorine Safety — pool type behaviour)",
     "⚠️ MINIMAL — Only 2/522 pools flagged. Extremely rare in dataset."),

    ("PISCINA COMUNITARIA", "Community Pool (flag)", "static_pool",
     "MEDIUM — Community pools typically have higher bather load and more chemical demand. Used in pipeline_v2 for pool_type derivation.",
     "P1 (Chlorine Safety — pool type), P4 (Chemical Cost — usage pattern)",
     "⚠️ PARTIAL — 174/522 pools (0.5% of rows). Static, 33% of fleet."),

    ("Piscina con skimmers", "Pool with Skimmers (count)", "static_pool",
     "LOW — Skimmer count indicates surface cleaning capacity. More skimmers = better organic removal.",
     "P2 (Filter Backwash — debris management context)",
     "⚠️ PARTIAL — 249/522 pools. Static, 48% of fleet."),

    ("Piscina desbordante", "Overflow Pool (flag)", "static_pool",
     "LOW — Overflow pools have continuous surface cleaning, different hydraulic behaviour.",
     "P1 (Chlorine Safety — pool type), P2 (Filter — hydraulic design)",
     "⚠️ MINIMAL — 43/522 pools. Static."),

    ("PISCINA EXTERIOR", "Outdoor Pool (flag)", "static_pool",
     "HIGH — Outdoor pools experience UV-driven chlorine degradation, rain dilution, and wind-borne debris. Used in pipeline_v2 for pool_type derivation.",
     "P1 (Chlorine Safety — sun exposure), P2 (Filter — contamination), P4 (Chemical Cost)",
     "⚠️ PARTIAL — 161/522 pools. Static, 31% of fleet."),

    ("Piscina ovalada", "Oval Pool (flag)", "static_pool",
     "LOW — Pool shape has minimal impact on water chemistry. Minor relevance to circulation patterns.",
     "None directly",
     "N/A — Marginal relevance"),

    ("PISCINA PARTICULAR", "Private Pool (flag)", "static_pool",
     "LOW — Private pools have lower bather load. Used in pipeline_v2 for pool_type derivation.",
     "P4 (Chemical Cost — lower demand)",
     "⚠️ MINIMAL — Only 8/522 pools."),

    ("PISCINA PUBLICA", "Public Pool (flag)", "static_pool",
     "MEDIUM — Public pools have highest bather load and strictest regulatory requirements.",
     "P1 (Chlorine Safety — high demand), P4 (Chemical Cost — high consumption)",
     "⚠️ MINIMAL — Only 1/522 pools."),

    ("(0714) Piscina rectangular", "Rectangular Pool type 0714 (flag)", "static_pool",
     "LOW — Pool shape variant. Minor relevance.",
     "None directly",
     "N/A — 111/522 pools"),

    ("(07) Piscina rectangular", "Rectangular Pool type 07 (flag)", "static_pool",
     "LOW — Pool shape variant. Minor relevance.",
     "None directly",
     "N/A — 74/522 pools"),

    ("Piscina redonda", "Round Pool (flag)", "static_pool",
     "LOW — Pool shape variant. Minor relevance.",
     "None directly",
     "N/A — 23/522 pools"),

    ("Superficie piscina", "Pool Surface Area (m²)", "static_pool",
     "HIGH — Larger surface area = more UV exposure = faster chlorine degradation. Also determines evaporation rate and debris accumulation. Used in Dose per m³ calculation.",
     "P1 (Chlorine Safety — UV surface), P4 (Chemical Cost — dose normalisation)",
     "⚠️ PARTIAL — 337/522 pools (1.0% of rows). Covers 65% of fleet. Static, backfillable."),

    ("VEGETACION CONTAMINANTE", "Nearby Vegetation (flag/level)", "static_pool",
     "MEDIUM — Vegetation drops organic matter (leaves, pollen) into the pool, increasing chlorine consumption and filter loading.",
     "P1 (Chlorine Safety — organic load), P2 (Filter Backwash — debris source)",
     "⚠️ MINIMAL — 47/522 pools (0.1% of rows). Only 9% of fleet."),

    ("Volumen piscina", "Pool Volume (m³)", "static_pool",
     "CRITICAL — Required for all chemical dosage calculations: kg_needed = concentration_deficit × volume × conversion_factor. Also used for Turnover Time and Dose per m³.",
     "P1 (Chlorine Safety — dosage), P4 (Chemical Cost — volume-normalised demand), P2 (Filter — turnover time)",
     "⚠️ PARTIAL — 351/522 pools (1.0% of rows). Covers 67% of fleet. Static, backfillable from company records."),

    ("Zona playa césped", "Grass Deck Area (flag)", "static_pool",
     "LOW — Grass surrounds introduce dirt and organic matter into the pool. Used in pipeline_v2 for deck_type derivation.",
     "P2 (Filter Backwash — contamination source)",
     "⚠️ PARTIAL — 107/522 pools. Static."),

    ("Zona PLAYA mixta", "Mixed Deck Area (flag)", "static_pool",
     "LOW — Mixed surrounds (grass + paving). Used in pipeline_v2 for deck_type derivation.",
     "P2 (Filter Backwash — contamination source)",
     "⚠️ PARTIAL — 82/522 pools. Static."),

    ("Zona PLAYA pavimentada", "Paved Deck Area (flag)", "static_pool",
     "LOW — Paved surrounds (cleanest). Used in pipeline_v2 for deck_type derivation.",
     "P2 (Filter Backwash — contamination source)",
     "⚠️ PARTIAL — 59/522 pools. Static."),

    ("Unnamed: 32", "— Blank separator —", "separator",
     "N/A — Empty column separating Sub-table 1 from Sub-table 2.",
     "None", "N/A"),

    # ── SUB-TABLE 2: Operations ──
    ("EMPLEADO.1", "Operations Technician", "meta",
     "LOW — Records who performed the operations check. Not used as ML feature.",
     "None",
     "N/A — 14.0% filled"),

    ("FECHA.1", "Operations Date", "meta",
     "IMPORTANT CONTEXT — The date when operations data was recorded. CRITICAL NOTE: This date does NOT align with the readings date (FECHA). Median gap is ~17.5 days. Operations data on each row was recorded on a different visit than the water quality reading.",
     "All operations-based predictions (alignment issue)",
     "⚠️ STRUCTURAL ISSUE — 15.4% filled, and misaligned with FECHA"),

    ("Horas dosificacion PH", "pH Dosing Hours", "dynamic_ops",
     "HIGH (for equipment prediction) — Hours the pH dosing pump ran. Required for Cumulative pH Dosing Hours calculation to predict pump wear.",
     "P5 (pH Dosing Pump Maintenance Cycle)",
     "❌ NO — 175 values (0.1%), only 57/522 pools, median 2/pool. Only 14 pools have ≥5 readings."),

    ("Horas filtracion diarias", "Daily Filtration Hours", "dynamic_ops",
     "HIGH — Hours the main pump ran per day. Required for Cumulative Pump Operating Hours to predict pump maintenance. Also indicates if operators are compensating for declining efficiency.",
     "P3 (Pump Degradation — runtime tracking), P5 (Main Pump Maintenance Cycle)",
     "⚠️ WEAK — 2,824 values (1.3%), 416/522 pools, but median only 6/pool. Only 2 pools have ≥3 consecutive pairs for trend computation."),

    ("Porcentaje dosificación PH", "pH Dosing Percentage (%)", "dynamic_ops",
     "HIGH (for equipment prediction) — Fraction of pH pump maximum capacity being used. Rising % with stable/worsening pH = pump wear signal.",
     "P3 (Pump Degradation — dosing effort), P5 (pH Dosing Pump — wear signal formula)",
     "❌ NO — 168 values (0.1%), 58/522 pools, median 2/pool. Only 10 pools have ≥5 readings."),

    ("Tiempo lavado /enjuague filtro", "Filter Wash / Rinse Time (min)", "dynamic_ops",
     "HIGH — Duration of each backwash cycle. Increasing wash time indicates filter media degradation. This is the strongest operational signal in the dataset.",
     "P2 (Filter Backwash — wash time trend), P5 (Filter Media Maintenance Cycle)",
     "⚠️ PARTIAL — 11,050 values (5.2%), 219/522 pools, median 44/pool. 139 pools have ≥3 consecutive pairs (2,378 total pairs). BEST dynamic operational field."),

    ("Horas dosificación hipo", "Chlorine Dosing Hours", "dynamic_ops",
     "HIGH (for equipment prediction) — Hours the chlorine dosing pump ran. Required for Cumulative Chlorine Dosing Hours.",
     "P3 (Pump Degradation), P5 (Chlorine Dosing Pump Maintenance Cycle)",
     "⚠️ WEAK — 4,427 values (2.1%), 282/522 pools, median 11/pool. 24 pools have ≥3 consecutive pairs. NOTE: NEVER co-occurs with Chlorine Dosing % on the same row (0 rows with both)."),

    ("Porcentaje dosificación hipoclorito", "Chlorine Dosing Percentage (%)", "dynamic_ops",
     "HIGH — Fraction of chlorine pump capacity being used. Rising % with declining chlorine = pump wear signal. Also indicates system effort to maintain chlorine levels.",
     "P1 (Chlorine Safety — dosing effort context), P3 (Pump Degradation), P5 (Chlorine Pump — wear signal)",
     "⚠️ PARTIAL — 6,130 values (2.9%), 285/522 pools, median 20/pool. 70 pools have ≥3 consecutive pairs (349 total). NOTE: NEVER co-occurs with Chlorine Dosing Hours (0 rows with both)."),

    ("Temperatura agua", "Water Temperature (°C)", "dynamic_ops",
     "HIGH — Higher temperatures accelerate chlorine decay and biological activity. The workflow document explicitly uses temperature for chlorine safety and chemical cost predictions.",
     "P1 (Chlorine Safety — temperature effect on decay), P4 (Chemical Cost — temperature-driven demand)",
     "⚠️ PARTIAL — 7,977 values (3.7%), 403/522 pools, median 16/pool. 94 pools with ≥3 consecutive pairs (467 total). Usable for a subset of pools."),

    ("Unnamed: 42", "— Blank separator —", "separator",
     "N/A — Empty column separating Sub-table 2 from Sub-table 3.",
     "None", "N/A"),

    # ── SUB-TABLE 3: Chemical Products Applied ──
    ("EMPLEADO.2", "Products Technician", "meta",
     "LOW — Records who applied the chemical products.",
     "None",
     "N/A — 55.1% filled"),

    ("FECHA.2", "Products Application Date", "meta",
     "MEDIUM — The date when chemicals were applied. Like FECHA.1, this does NOT align with the readings date. Used in pipeline_v2 via merge_asof temporal join (within 14-day window).",
     "P4 (Chemical Cost — timing of chemical application)",
     "⚠️ STRUCTURAL ISSUE — 56.9% filled, misaligned with FECHA"),

    ("T-500 (GRUPO QP)", "T-500 Chlorine Product (kg)", "product",
     "MEDIUM — A specific chlorine product (brand: Grupo QP). Quantity applied. Aggregated into total_chlorine_applied in pipeline_v2.",
     "P4 (Chemical Cost — historical consumption)",
     "✅ GOOD — 33,688 values (15.8%), 461/522 pools"),

    ("ALBORAL TABLETAS 250 GRS RF. 201710", "Alboral Chlorine Tablets (kg)", "product",
     "LOW — A specific chlorine tablet brand. Aggregated into total_chlorine_applied.",
     "P4 (Chemical Cost)",
     "⚠️ SPARSE — 767 values (0.4%), 223 pools"),

    ("FLOVIL PASTILLAS", "Flovil Flocculant Tablets (units)", "product",
     "MEDIUM — Flocculant product for water clarity. Indicates turbidity management activity.",
     "P2 (Filter Backwash — turbidity treatment), P4 (Chemical Cost)",
     "✅ GOOD — 21,001 values (9.9%), 462 pools"),

    ("HIPO GARRAFAS 20KG.", "Hypochlorite Jugs 20kg (units)", "product",
     "MEDIUM — Liquid chlorine in large containers. Aggregated into total_chlorine_applied.",
     "P1 (Chlorine Safety — dosing response), P4 (Chemical Cost)",
     "⚠️ PARTIAL — 4,575 values (2.1%), 369 pools"),

    ("HIPO GR CHLORYTE", "Chloryte Granular Hypochlorite (kg)", "product",
     "MEDIUM — Granular chlorine product. Aggregated into total_chlorine_applied.",
     "P1 (Chlorine Safety), P4 (Chemical Cost)",
     "⚠️ PARTIAL — 7,071 values (3.3%), 459 pools"),

    ("HIPO GRANULADO XAKA", "Xaka Granular Hypochlorite (kg)", "product",
     "MEDIUM — Granular chlorine product (Xaka brand). Aggregated into total_chlorine_applied.",
     "P1 (Chlorine Safety), P4 (Chemical Cost)",
     "⚠️ PARTIAL — 3,302 values (1.6%), 396 pools"),

    ("HIPO STICKS BAYROL", "Bayrol Hypochlorite Sticks (kg)", "product",
     "LOW — Chlorine sticks (Bayrol brand). Aggregated into total_chlorine_applied.",
     "P1 (Chlorine Safety), P4 (Chemical Cost)",
     "⚠️ SPARSE — 3,083 values (1.4%), 210 pools"),

    ("HIPO TAB. RITOCAL", "Ritocal Hypochlorite Tablets (kg)", "product",
     "LOW — Chlorine tablets (Ritocal brand). Aggregated into total_chlorine_applied.",
     "P4 (Chemical Cost)",
     "⚠️ SPARSE — 4,268 values (2.0%), 345 pools"),

    ("HIPO TABLETAS 200Gr. QP", "QP 200g Hypochlorite Tablets (kg)", "product",
     "LOW — Chlorine tablets (QP brand). Aggregated into total_chlorine_applied.",
     "P4 (Chemical Cost)",
     "⚠️ SPARSE — 2,454 values (1.2%), 273 pools"),

    ("HIPO TABLETAS XAKA", "Xaka Hypochlorite Tablets (kg)", "product",
     "MEDIUM — Chlorine tablets (Xaka brand). Aggregated into total_chlorine_applied.",
     "P1 (Chlorine Safety), P4 (Chemical Cost)",
     "⚠️ PARTIAL — 8,774 values (4.1%), 384 pools"),

    ("PH MINUS GRANULADO 6kg", "pH Minus Granular 6kg (units)", "product",
     "MEDIUM — pH-lowering product (granular form). Aggregated into total_ph_minus_product in pipeline_v2.",
     "P4 (Chemical Cost — pH correction demand)",
     "⚠️ PARTIAL — 3,171 values (1.5%), 343 pools"),

    ("PH MINUS LIQUIDO 13.5 KG", "pH Minus Liquid 13.5kg (units)", "product",
     "MEDIUM — pH-lowering product (liquid form). Aggregated into total_ph_minus_product.",
     "P4 (Chemical Cost — pH correction demand)",
     "⚠️ PARTIAL — 4,651 values (2.2%), 403 pools"),

    ("PH MINUS LIQUIDO 27 KG.", "pH Minus Liquid 27kg (units)", "product",
     "LOW — pH-lowering product (large container). Aggregated into total_ph_minus_product.",
     "P4 (Chemical Cost)",
     "⚠️ SPARSE — 369 values (0.2%), 87 pools"),

    ("PROTECT & SHINE", "Protect & Shine (units)", "product",
     "LOW — Pool surface treatment product. Not directly related to water chemistry predictions.",
     "P4 (Chemical Cost — minor)",
     "⚠️ SPARSE — 516 values (0.2%), 88 pools"),

    ("SG XAKA (AGONET GR90)", "SG Xaka Algaecide/Clarifier (kg)", "product",
     "MEDIUM — Algaecide/clarifier product. Indicates pools with algae problems or preventive treatment.",
     "P2 (Filter — water clarity), P4 (Chemical Cost)",
     "⚠️ PARTIAL — 10,658 values (5.0%), 435 pools"),

    ("SUPERKLAR", "Superklar Flocculant/Clarifier (units)", "product",
     "MEDIUM — Flocculant/clarifier product. Indicates turbidity management.",
     "P2 (Filter — water clarity), P4 (Chemical Cost)",
     "✅ GOOD — 12,753 values (6.0%), 470 pools"),
]


# ─────────────────────────────────────────────────
# Build the DOCX
# ─────────────────────────────────────────────────
doc = Document()

# --- Page setup: landscape for wide table ---
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)  # A4 width
    section.page_height = Cm(21.0) # A4 height
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title ---
title = doc.add_heading('Data Column Feasibility Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Smart Predictive Pool Maintenance System — SPP System (Pepe Gutiérrez)')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

# --- Report metadata ---
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta_p.add_run(
    f'Dataset: merged_pool_data_2017_2022.csv  |  '
    f'{total_rows:,} rows  |  {total_pools} pools  |  '
    f'{date_min.strftime("%Y-%m-%d")} to {date_max.strftime("%Y-%m-%d")}  |  '
    f'Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}'
)
meta_run.font.size = Pt(10)
meta_run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_paragraph()

# --- Prediction Legend ---
doc.add_heading('Prediction Reference (from Factors Workflow Document)', level=1)

legend_data = [
    ("P1", "Chlorine Safety Forecast", "Predict abnormal chlorine decay and generate safety alerts"),
    ("P2", "Filter Backwash & Replacement", "Predict filter degradation, backwash timing, and replacement needs"),
    ("P3", "Pump Degradation", "Predict pump wear and maintenance needs before failure"),
    ("P4", "Chemical Cost Forecast", "Forecast future chemical demand and operational costs"),
    ("P5", "Equipment Maintenance Cycle", "Track cumulative hours and performance trends for all equipment"),
]

legend_table = doc.add_table(rows=1, cols=3)
legend_table.style = 'Light List Accent 1'
legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = legend_table.rows[0].cells
hdr[0].text = 'Code'
hdr[1].text = 'Prediction Name'
hdr[2].text = 'Description'

for code, name, desc in legend_data:
    row = legend_table.add_row().cells
    row[0].text = code
    row[1].text = name
    row[2].text = desc

doc.add_paragraph()

# --- Pipeline Status ---
doc.add_heading('Current Pipeline Implementation Status (pipeline_v2.py)', level=1)

status_data = [
    ("P1 — Chlorine Safety Forecast", "~85%",
     "Implemented. Predicts next-visit chlorine level using decay rate, pH interaction, lags, rolling stats, and headroom. Generates URGENT alerts when chlorine < 0.5 or > 5.0 mg/L. Missing: temperature integration (data sparse), pH-Chlorine Effectiveness Index (not hardcoded but XGBoost learns interaction)."),
    ("P2 — Filter Backwash & Replacement", "~40%",
     "Partially implemented. Predicts next-visit turbidity and prescribes flocculant. Missing: pump flow rate trends, hydraulic efficiency, wash time trends, turnover time — all blocked by sparse operations data."),
    ("P3 — Pump Degradation", "0%",
     "Not implemented. Requires cumulative operating hours, flow rate tracking, and dosing percentage trends — all critically sparse in the dataset."),
    ("P4 — Chemical Cost Forecast", "~75%",
     "Substantially implemented. Predicts chemical dosage in kg using water quality forecasts + volume-based mass-balance formulas. Missing: Euro (€) cost conversion, temperature-driven demand adjustment."),
    ("P5 — Equipment Maintenance Cycle", "0%",
     "Not implemented. Requires consistent recording of dosing hours, dosing percentages, pump flow rates, and filter wash times at every visit. Data is too sparse and the required field pairs never co-occur."),
]

status_table = doc.add_table(rows=1, cols=3)
status_table.style = 'Light List Accent 1'
status_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = status_table.rows[0].cells
hdr[0].text = 'Prediction'
hdr[1].text = 'Implementation'
hdr[2].text = 'Details'

for pred, pct, details in status_data:
    row = status_table.add_row().cells
    row[0].text = pred
    row[1].text = pct
    row[2].text = details

doc.add_paragraph()

# --- MAIN TABLE ---
doc.add_heading('Complete Column Mapping and Data Availability', level=1)

note = doc.add_paragraph()
note_run = note.add_run(
    'Note on Data Availability Interpretation: Static fields (pool-level constants like volume, filter diameter, pump flow rate) '
    'only need one value per pool — a low row percentage is expected and not a problem if the pool count is reasonable. '
    'Dynamic fields (recorded at each visit like dosing hours, wash time, temperature) need many values per pool '
    'to compute trends. The "Can We Build It?" column considers both the field type and the per-pool density.'
)
note_run.font.size = Pt(10)
note_run.font.italic = True
note_run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Create table
headers = [
    'Spanish Column Name',
    'English Translation',
    'Field Type',
    'ML Relevance & Role',
    'Predictions Affected',
    'Data Availability\n(merged_pool_data\n_2017_2022.csv)',
    'Can We Build the\nPrediction?'
]

table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

# Header row
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for paragraph in hdr_cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)

# Populate data rows
for spanish, english, field_type, relevance, predictions, can_build in columns_info:
    # Skip separators in visual display but include them
    nn = non_null(spanish)
    pct = fill_pct(spanish)
    pc = pool_count(spanish)

    if field_type == "separator":
        availability = "Empty column (0 values)"
    elif field_type in ("static_pool", "static_equipment"):
        availability = f"{nn:,} values ({pct}% of rows)\n{pc}/{total_pools} pools\n[Static — 1 value per pool sufficient]"
    elif field_type == "dynamic_ops":
        mn, md, mx, gte10 = per_pool_stats(spanish)
        availability = f"{nn:,} values ({pct}% of rows)\n{pc}/{total_pools} pools\nPer-pool: min={mn}, med={md}, max={mx}\nPools with ≥10 readings: {gte10}"
    elif field_type == "product":
        availability = f"{nn:,} values ({pct}% of rows)\n{pc}/{total_pools} pools"
    elif field_type == "core_reading":
        mn, md, mx, gte10 = per_pool_stats(spanish)
        availability = f"{nn:,} values ({pct}% of rows)\n{pc}/{total_pools} pools\nPer-pool: min={mn}, med={md}, max={mx}\nPools with ≥10: {gte10}"
    else:  # meta
        availability = f"{nn:,} values ({pct}% of rows)\n{pc}/{total_pools} pools"

    # Field type label
    type_labels = {
        "core_reading": "Core Water Reading",
        "static_pool": "Static — Pool Feature",
        "static_equipment": "Static — Equipment Spec",
        "dynamic_ops": "Dynamic — Operations",
        "product": "Chemical Product",
        "meta": "Metadata / ID",
        "separator": "— Separator —",
    }

    row = table.add_row().cells
    row[0].text = spanish
    row[1].text = english
    row[2].text = type_labels.get(field_type, field_type)
    row[3].text = relevance
    row[4].text = predictions
    row[5].text = availability
    row[6].text = can_build

    # Font size for all cells
    for cell in row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

# Set column widths
widths = [Cm(3.5), Cm(3.5), Cm(2.0), Cm(7.0), Cm(3.5), Cm(4.0), Cm(3.5)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width


# --- Critical Findings Section ---
doc.add_page_break()
doc.add_heading('Critical Findings and Recommendations', level=1)

doc.add_heading('1. Dataset Scale Improvement', level=2)
p = doc.add_paragraph()
p.add_run('The new merged dataset (2017–2022) is a ').font.size = Pt(11)
r = p.add_run('massive improvement')
r.bold = True
r.font.size = Pt(11)
p.add_run(' over the original 2022-only dataset:').font.size = Pt(11)

scale_table = doc.add_table(rows=1, cols=3)
scale_table.style = 'Light List Accent 1'
hdr = scale_table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Old Dataset (2022 only)'
hdr[2].text = 'New Dataset (2017–2022)'

scale_data = [
    ("Total rows", "4,231", f"{total_rows:,}"),
    ("Unique pools", "46", f"{total_pools}"),
    ("Date range", "Jan–Dec 2022 (1 year)", "Apr 2017 – Dec 2022 (5.5 years)"),
    ("pH readings", "3,719", f"{non_null('PH'):,}"),
    ("Free chlorine readings", "3,718", f"{non_null('CLORO LIBRE'):,}"),
    ("Turbidity readings", "3,632", f"{non_null('TURBIDEZ'):,}"),
    ("Filter wash time values", "381", f"{non_null('Tiempo lavado /enjuague filtro'):,}"),
    ("Chlorine dosing % values", "455", f"{non_null('Porcentaje dosificación hipoclorito'):,}"),
    ("Water temperature values", "854", f"{non_null('Temperatura agua'):,}"),
]

for metric, old, new in scale_data:
    row = scale_table.add_row().cells
    row[0].text = metric
    row[1].text = old
    row[2].text = new

doc.add_paragraph()

doc.add_heading('2. What Can Be Built NOW (with this dataset)', level=2)
p = doc.add_paragraph()
p.add_run('Water Chemistry Predictions (P1 + P4): ').bold = True
p.add_run(
    'Fully buildable and significantly improved. With 195,000+ pH readings, 191,000+ chlorine readings, '
    'and 150,000+ turbidity readings across 480+ pools over 5.5 years, the XGBoost models will have '
    'dramatically better training data. Multi-year seasonal patterns (summer vs winter) can now be learned properly. '
    'The chemical product data (Sub-table 3) is well-populated for cost forecasting.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Filter Wash Time Monitoring (partial P2 + P5): ').bold = True
p.add_run(
    'Now partially buildable. 11,050 wash time values across 219 pools (median 44 readings/pool) '
    'with 139 pools having ≥3 consecutive pairs (2,378 total pairs). This is sufficient to build a '
    'wash time trend monitor that detects filter degradation for a meaningful subset of the fleet.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Chlorine Dosing % Trend (partial P3 + P5): ').bold = True
p.add_run(
    'Marginally buildable. 6,130 values across 285 pools with 70 pools having ≥3 consecutive pairs. '
    'A simplified chlorine pump wear signal (tracking rising dosing %) could be implemented for these 70 pools, '
    'though without co-occurring dosing hours data, the full prediction logic from the workflow document cannot be replicated.'
)

doc.add_paragraph()

doc.add_heading('3. What CANNOT Be Built (data blockers)', level=2)

p = doc.add_paragraph()
p.add_run('pH Dosing Pump Prediction (P5): ').bold = True
p.add_run(
    'Still impossible. Only 175 pH dosing hour values and 168 pH dosing % values across the entire 5.5-year dataset. '
    'This data was simply never recorded consistently.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Main Pump Degradation (P3 + P5): ').bold = True
p.add_run(
    'Still blocked. Daily filtration hours have only 2 pools with ≥3 consecutive readings for trend computation. '
    'The main pump flow rate is a static field (254/522 pools) which is adequate, but without temporal filtration hour data, '
    'cumulative tracking and hydraulic efficiency trending cannot be computed.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Cross-field co-occurrence problem: ').bold = True
p.add_run(
    'Chlorine dosing hours and chlorine dosing % NEVER appear on the same row (0 co-occurrences across 212,850 rows). '
    'Same for pH dosing hours and pH dosing %. The workflow document requires combining these signals for pump '
    'degradation confirmation — this is structurally impossible with the current data collection process.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Sub-table date misalignment: ').bold = True
p.add_run(
    'The three sub-tables (readings, operations, products) use independent date columns that do not correspond to '
    'the same technician visit. The operations date (FECHA.1) is typically 17+ days offset from the readings date (FECHA). '
    'This means operational data cannot be reliably correlated with the water quality reading on the same row.'
)

doc.add_paragraph()

doc.add_heading('4. Recommended Actions for Data Collection', level=2)
p = doc.add_paragraph(
    'To enable the full Smart Predictive Pool Maintenance System as described in the Factors Workflow document, '
    'the following data collection improvements are recommended for the company\'s field technicians:'
)

actions = [
    ("Mandatory daily filtration hours recording", "Record at EVERY visit, not sporadically. This unlocks main pump cumulative tracking."),
    ("Mandatory chlorine and pH dosing hours + percentage", "Record BOTH fields at the SAME visit. Currently they never co-occur, preventing pump wear confirmation."),
    ("Align operations data with readings", "Operations measurements should be recorded on the SAME visit as the water quality reading, using the SAME date stamp."),
    ("Backfill static equipment specs", "Pool volume, surface area, filter diameter, number of filters/motors, and pump flow rates should be recorded for all 522 pools. These are one-time entries."),
    ("Record water temperature at every visit", "Temperature is critical for chlorine decay modelling. Currently only 3.7% of visits have it."),
]

for title_text, desc in actions:
    p = doc.add_paragraph()
    p.add_run(f'{title_text}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# --- Footer ---
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run(
    f'Report generated on {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")} | '
    f'Dataset: merged_pool_data_2017_2022.csv | '
    f'Reference: Factors workflow Smart Predictive Pool Maintenance System.docx | '
    f'Pipeline: pipeline_v2.py'
)
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(150, 150, 150)
r.font.italic = True

# --- Save ---
output_path = "outputs/data_feasibility_report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"Total columns mapped: {len(columns_info)}")
