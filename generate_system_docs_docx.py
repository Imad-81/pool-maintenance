"""
Script to generate the comprehensive, professional .docx documentation for the
Spain (Alicante) Pool Predictive Maintenance and Operations System (V6.0).
"""

import os
import json
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

PROJECT_ROOT = Path(__file__).resolve().parent
DOCX_OUTPUT_PATH = PROJECT_ROOT / "Spain_Pool_Predictive_Maintenance_Complete_System_Documentation_V6.docx"

# Color Palette Constants
HEX_NAVY = "0F2C59"       # Primary Headers & Accents
HEX_OCEAN = "3085C3"      # Subheaders & Table Headers
HEX_LIGHT_BLUE = "EAF2F8" # Table Header Alt / Light Background
HEX_CHARCOAL = "1E293B"   # Body Text
HEX_MUTED = "64748B"      # Subtitles & Captions
HEX_BORDER = "CBD5E1"     # Table Borders
HEX_ZEBRA = "F8FAFC"      # Alternating Table Row Shading
HEX_ALERT_RED = "DC2626"  # Red alerts
HEX_ALERT_AMBER = "D97706"# Amber alerts
HEX_ALERT_GREEN = "16A34A"# Green alerts
HEX_CODE_BG = "F1F5F9"    # Code background

COLOR_NAVY = RGBColor(15, 44, 89)
COLOR_OCEAN = RGBColor(48, 133, 195)
COLOR_CHARCOAL = RGBColor(30, 41, 59)
COLOR_MUTED = RGBColor(100, 116, 139)
COLOR_RED = RGBColor(220, 38, 38)
COLOR_AMBER = RGBColor(217, 119, 6)
COLOR_GREEN = RGBColor(22, 163, 74)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_MUTED
    return p


def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_NAVY
    return p


def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_OCEAN
    return p


def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_CHARCOAL
    return p


def add_paragraph(doc, text="", bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_CHARCOAL
    if text:
        r_text = p.add_run(text)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(10.5)
        r_text.font.italic = italic
        r_text.font.color.rgb = COLOR_CHARCOAL
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(10.5)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_CHARCOAL
    if text:
        r_text = p.add_run(text)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = COLOR_CHARCOAL
    return p


def add_callout(doc, title, text, callout_type="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    
    border_color = HEX_OCEAN
    bg_color = "F0F7FF"
    if callout_type == "WARNING":
        border_color = HEX_ALERT_AMBER
        bg_color = "FFFBEB"
    elif callout_type == "IMPORTANT" or callout_type == "CAUTION":
        border_color = HEX_ALERT_RED
        bg_color = "FEF2F2"
    elif callout_type == "SUCCESS":
        border_color = HEX_ALERT_GREEN
        bg_color = "F0FDF4"

    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    prefix = f"[{callout_type}] " if callout_type != "NOTE" else ""
    r_title = p.add_run(f"{prefix}{title}\n")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    if callout_type in ("WARNING", "IMPORTANT", "CAUTION"):
        r_title.font.color.rgb = COLOR_RED if callout_type in ("IMPORTANT", "CAUTION") else COLOR_AMBER
    else:
        r_title.font.color.rgb = COLOR_NAVY

    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(10.0)
    r_text.font.color.rgb = COLOR_CHARCOAL

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)


def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, HEX_CODE_BG)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(9.0)
    run.font.color.rgb = COLOR_CHARCOAL

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(4)


def add_styled_table(doc, headers, data, col_widths=None):
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl)

    # Header Row
    hdr_row = tbl.rows[0]
    hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    for idx, heading in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_shading(cell, HEX_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(str(heading))
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row = tbl.rows[r_idx + 1]
        bg_color = HEX_ZEBRA if (r_idx % 2 == 1) else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(9.0)
            run.font.color.rgb = COLOR_CHARCOAL

    if col_widths and len(col_widths) == len(headers):
        for row in tbl.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)
    return tbl


def add_image_with_caption(doc, image_path, caption, width=Inches(6.0)):
    if not Path(image_path).exists():
        add_paragraph(doc, f"[Image file not found at {image_path}]", italic=True)
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    p_img.paragraph_format.keep_with_next = True
    run_img = p_img.add_run()
    run_img.add_picture(str(image_path), width=width)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after = Pt(10)
    run_cap = p_cap.add_run(caption)
    run_cap.font.name = "Calibri"
    run_cap.font.size = Pt(9.0)
    run_cap.font.italic = True
    run_cap.font.color.rgb = COLOR_MUTED


def build_complete_document():
    doc = Document()

    # Set Standard Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # -----------------------------------------------------------------------
    # COVER / TITLE BLOCK
    # -----------------------------------------------------------------------
    add_title(doc, "SPAIN (ALICANTE) COLLECTIVE-USE POOLS")
    add_subtitle(doc, "Predictive Maintenance, Physical Kinetics Rate Integration, Automated Chemical Dosing, and Fleet Operations System (Version 6.0)")

    meta_headers = ["Document Attribute", "Specification & Implementation Details"]
    meta_data = [
        ["System Version", "Version 6.0 (Post-Treatment Setpoint Re-Anchor & Hybrid Physics Architecture)"],
        ["Target Region & Coordinates", "Alicante, Spain (Latitude: 38.3452° N, Longitude: -0.4815° W)"],
        ["Regulatory Framework", "Real Decreto 742/2013 (National) & Decreto 85/2018 (Comunitat Valenciana)"],
        ["Core Technology Stack", "Python 3.10+, FastAPI, XGBoost 3.2+, PostgreSQL 16, Prisma ORM, React 19, Vite"],
        ["Dataset Baseline", "42,617 rows (2023–2026) across 135 liquid chlorine dosing pump community pools"],
        ["External Intelligence", "Open-Meteo High-Resolution UV, Solar Radiation, Temperature & Wind APIs"],
        ["Operational Status", "Production-Ready with Automated APScheduler Cron Jobs & Web Dashboard"],
    ]
    add_styled_table(doc, meta_headers, meta_data, [Inches(2.2), Inches(4.3)])

    add_callout(
        doc,
        "EXECUTIVE BRIEF FOR ENGINEERING LEADERSHIP & CLIENT OPERATORS",
        "This document provides the exhaustive, end-to-end technical specification of the Pool Predictive Maintenance "
        "System (V6.0). It details every layer of the architecture: raw data ingestion and Spanish header normalization, "
        "the client-directed liquid chlorine pump scoping, the physical-kinetics rate integration engine, the 87-feature ML "
        "pipeline, the XGBoost next-day regressors, the multi-day chained forecast algorithm, the 525-grid chemical dosing "
        "optimizer, the PostgreSQL/Prisma relational schema, the FastAPI REST API, and the React 19 operator dashboard.",
        "NOTE"
    )

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # CHAPTER 1: PROBLEM DOMAIN & REGULATORY GROUNDING
    # -----------------------------------------------------------------------
    add_heading_1(doc, "1. Problem Domain, Regulatory Grounding & Spanish Mediterranean Realities")

    add_heading_2(doc, "1.1 The Operational Challenge in Collective-Use Pools")
    add_paragraph(doc, "In the province of Alicante, Spain, collective-use swimming pools (piscinas de uso colectivo)—installed in residential urbanisations, resort complexes, and hotels—operate under stringent legal requirements and severe environmental stress. Maintaining water quality is critical: inadequate disinfectant levels expose hundreds of daily bathers to microbiological pathogens (such as Pseudomonas aeruginosa and Escherichia coli), while excessive chemical concentrations cause toxic eye, mucosal, and dermal burns.")

    add_paragraph(doc, "Historically, pool maintenance companies dispatched technicians every 2 to 4 days to perform manual water testing, log readings, and adjust chemical feeder pumps. However, this reactive workflow has severe limitations:")
    add_bullet(doc, "Technicians only observe the water state at the exact moment of their visit, leaving 48 to 96 hours of unmonitored degradation between visits.", "Inter-Visit Blindspots: ")
    add_bullet(doc, "Alicante experiences extreme summer solar radiation (>30 MJ/m²) and UV indices (>9.5), causing rapid hypochlorous acid photolysis that can deplete a safe chlorine buffer in hours.", "Atmospheric Acceleration: ")
    add_bullet(doc, "Manual rule-of-thumb pump adjustments lead to erratic swings between under-chlorination hazards and severe chemical over-dosing.", "Heuristic Dosing Inefficiencies: ")

    add_heading_2(doc, "1.2 Spanish Regulatory Framework (Real Decreto 742/2013 & Decreto 85/2018)")
    add_paragraph(doc, "The system's safety boundaries, headroom calculations, and alarm states are strictly anchored in Spanish national and regional legislation:")
    add_bullet(doc, "Establishes mandatory water quality criteria, autocontrol protocols, and corrective action standards for all collective-use pools in Spain.", "Real Decreto 742/2013 (National Spanish Standard): ")
    add_bullet(doc, "Governs the obligatory daily autocontrol logbook (Libro de Registro de Control del Agua) and inspection criteria in the Valencian Community.", "Decreto 85/2018 (Comunitat Valenciana): ")

    reg_headers = ["Parameter", "RD 742/2013 Regulatory Range", "Client Optimal Target", "Hazard & Breach Actions"]
    reg_data = [
        ["Free Chlorine (Cloro Libre)", "0.5 – 2.0 mg/L (ideal upper)", "1.0 – 1.5 mg/L (ideal 1.25)", "< 0.5 mg/L: Pathogen risk (Immediate Visit 🚨)\n> 5.0 mg/L: Mandatory pool closure (RD 742/2013)"],
        ["pH", "7.2 – 8.0 pH units", "7.2 – 7.8 (ideal 7.4)", "< 7.2: Eye/skin irritation, pipe corrosion 🚨\n> 8.0: Scale formation, chlorine inactivation 🚨"],
        ["Turbidity (Turbidez)", "≤ 5.0 NTU (limit)", "≤ 1.0 NTU (ideal ≤ 0.5)", "> 5.0 NTU: Cloudiness, filtration failure (Immediate Visit 🚨)"],
    ]
    add_styled_table(doc, reg_headers, reg_data, [Inches(1.8), Inches(1.7), Inches(1.5), Inches(1.5)])

    add_heading_2(doc, "1.3 The Mediterranean '60% Chlorine Overdose' Phenomenon")
    add_paragraph(doc, "A critical discovery made during field data validation with Iberpiscinas SLU technicians (led by Jesús Santana) is the intentional Mediterranean Overdosing practice. Due to intense Alicante solar radiation and sudden weekend bather spikes, technicians routinely adjust liquid chlorine dosing pumps to achieve 2.0 to 4.0 mg/L free chlorine upon departure. While 2.0 mg/L is the nominal ideal ceiling under RD 742/2013 Annexe I, Spanish law only mandates pool closure if free chlorine exceeds 5.0 mg/L. The V6 system explicitly accounts for this field practice: levels between 2.0 and 4.0 mg/L are classified as 'Spanish Mediterranean Buffer (Monitor)' rather than immediate emergencies, while optimizing dosing toward the client's sustainable 1.0–1.5 mg/L ideal.")

    add_heading_2(doc, "1.4 Pre-Treatment Measurement Reality vs. Post-Treatment Setpoint Assumption")
    add_callout(
        doc,
        "CORE METHODOLOGICAL INNOVATION: THE POST-TREATMENT SETPOINT RE-ANCHOR",
        "Field inspection records in commercial pool maintenance reflect PRE-TREATMENT water states. When a technician arrives, "
        "they measure the degraded water resulting from the previous inter-visit gap, record the numbers, and then add chemicals or "
        "adjust pumps to restore ideal balance. Water degradation therefore evolves FROM the post-treatment state, not from the "
        "degraded pre-treatment reading.\n\n"
        "In V6.0, the pipeline anchors synthetic targets and inference kinetics to a configurable post-treatment setpoint "
        "(Free Cl: 2.5 mg/L, pH: 7.4, Turbidity: 0.5 NTU). This eliminated artificial R² inflation from lag-copying and lowered "
        "true Free Chlorine Mean Absolute Error (MAE) from 0.26 to 0.1972 mg/L.",
        "IMPORTANT"
    )

    # -----------------------------------------------------------------------
    # CHAPTER 2: RAW DATASET, SCOPING & DATA CLEANING
    # -----------------------------------------------------------------------
    add_heading_1(doc, "2. Dataset Architecture, Target Fleet Scoping & Data Ingestion")

    add_heading_2(doc, "2.1 Raw Dataset Composition")
    add_paragraph(doc, "The raw operational data originated from Pepe Gutiérrez's SPP System maintenance database in Alicante, consolidated into 'data/Merged_2023_2026.xlsx'. The dataset spans January 2, 2023 through August 5, 2026, comprising 42,617 raw rows across 61 denormalized columns.")

    add_paragraph(doc, "Each row in the master sheet represents three operational sub-tables logged side-by-side:")
    add_bullet(doc, "Pool ID, community name, timestamp, technician, measured pH, free chlorine, turbidity, and physical dimensions.", "Sub-Table 1 (Water Quality Readings): ")
    add_bullet(doc, "Dosing pump hours, filtration run hours, hypochlorite dosing %, pH dosing %, and water temperature.", "Sub-Table 2 (Equipment Operations): ")
    add_bullet(doc, "Granular/liquid hypochlorite (kg), pH minus liquid/granules (kg), and clarifiers/flocculants (tablets/sticks).", "Sub-Table 3 (Applied Chemical Products): ")

    add_heading_2(doc, "2.2 Liquid Chlorine Dosing Pump Scoping")
    add_paragraph(doc, "Per client requirements, predictive modeling is scoped exclusively to community pools equipped with automated liquid chlorine dosing pumps. The reference registry 'data/Listado_piscinas_bomba_cloro.xlsx' lists 138 qualifying installations.")
    add_bullet(doc, "126 pools were matched unambiguously by extracting the numeric reference enclosed in parentheses (e.g., 'Cabo Verde (19)' -> '19').", "Primary Exact Reference Match: ")
    add_bullet(doc, "9 compound or multi-pool community entries (e.g., '654-655') were resolved through normalized substring community reconciliation.", "Secondary Fuzzy Community Match: ")
    add_bullet(doc, "135 active qualifying pools representing 38,362 validated reading rows are retained in the master dataset. Non-qualifying pools (~4,255 rows) are safely filtered out.", "Final Scoped Fleet: ")

    add_heading_2(doc, "2.3 Column Normalization & Multi-Visit Deduplication")
    add_paragraph(doc, "1. Spanish Column Header Mapping: The pipeline maps 56 raw Spanish column names into standardized snake_case identifiers using a static dictionary (RENAME_MAP) in 'ml/config.py', ensuring total resilience against column reordering in Excel exports.")
    add_paragraph(doc, "2. Multi-Visit Day Deduplication: When multiple visits occur on the same calendar day for a single pool (1,061 historical instances), the pipeline preserves the LAST visit of the day as the official end-of-day pool state. To preserve operational incident history, a binary flag 'multi_visit_day = 1' is generated and passed into the feature engine.")

    add_heading_2(doc, "2.4 Static Pool Metadata Backfilling & Fleet Imputation")
    add_paragraph(doc, "Physical pool characteristics (volume m³, surface area m², filter diameter mm, filter count, motor count) suffered from >50% missingness in raw logging. Because physical dimensions are static properties of each pool, the pipeline implements a two-stage backfilling strategy:")
    add_bullet(doc, "Extracts the maximum non-null recorded dimension for that specific pool ID across all historical records.", "Stage 1 (Per-Pool Forward/Backward Fill): ")
    add_bullet(doc, "For pools lacking any dimension record, backfills the fleet-wide median (median volume: 225.0 m³, median surface area: 157.5 m², filter diameter: 900.0 mm).", "Stage 2 (Fleet-Wide Median Imputation): ")
    add_paragraph(doc, "This procedure elevated static feature completeness from 1.1% to 100%, unlocking crucial volume-normalized dosage features.")

    add_heading_2(doc, "2.5 As-Of Backward Merging of Operations & Products")
    add_paragraph(doc, "Because equipment operations and chemical applications were occasionally logged asynchronously from water quality measurements, the pipeline executes a pool-by-pool backward as-of merge (pd.merge_asof) with a 14-day tolerance window. Each water quality reading is matched to the most recent preceding equipment configuration and product application.")

    # -----------------------------------------------------------------------
    # CHAPTER 3: WEATHER INTELLIGENCE
    # -----------------------------------------------------------------------
    add_heading_1(doc, "3. External Weather Intelligence & Open-Meteo Integration")

    add_paragraph(doc, "Chlorine degradation in outdoor pools is predominantly driven by solar ultraviolet radiation and elevated ambient temperatures. The system integrates high-resolution daily atmospheric intelligence for Alicante (38.3452° N, -0.4815° W) fetched from the Open-Meteo European Weather Model API and cached in 'data/weather_alicante_2023_2026.csv' (1,312 continuous days).")

    wx_headers = ["Weather Signal Group", "Features Included", "Physical Mechanism Captured"]
    wx_data = [
        ["Current Day Weather (9 features)", "w_temp_max, w_temp_mean, w_uv_max, w_uv_clear_sky_max, w_solar_radiation, w_sunshine_hours, w_precipitation_mm, w_wind_max_kmh, w_et0", "Captures immediate UV photolysis rates, evaporation loss, thermal acceleration of oxidation, and wind-borne debris."],
        ["Cumulative Weather Since Last Visit (4 features)", "w_uv_sum_since, w_solar_sum_since, w_precip_sum_since, w_temp_mean_since", "Integrates the total atmospheric energy and rainfall accumulated over the k-day inter-visit gap."],
        ["Tomorrow Prediction-Day Forecast (9 features)", "w_tmrw_temp_max, w_tmrw_temp_mean, w_tmrw_uv_max, w_tmrw_uv_clear_sky_max, w_tmrw_solar_radiation, w_tmrw_sunshine_hours, w_tmrw_precipitation_mm, w_tmrw_wind_max_kmh, w_tmrw_et0", "Provides direct forward-looking signals for next-day chemical demand and impending weather shocks."],
    ]
    add_styled_table(doc, wx_headers, wx_data, [Inches(2.0), Inches(2.2), Inches(2.3)])

    add_paragraph(doc, "Merge Integrity Enforcement: Weather joins are executed via exact date left joins on normalized timestamps. Strict runtime assertions ensure zero row-count inflation (38,362 -> 38,362 rows).")

    # -----------------------------------------------------------------------
    # CHAPTER 4: FEATURE ENGINEERING
    # -----------------------------------------------------------------------
    add_heading_1(doc, "4. Complete Feature Engineering Pipeline (87 Numeric Signals)")

    add_paragraph(doc, "The pipeline computes 87 numeric and categorical features structured into 13 distinct functional groups. Strict bit-for-bit feature parity is maintained between training and inference:")

    feat_headers = ["Group", "Feature Names", "Mathematical Formulation & Meaning"]
    feat_data = [
        ["Static Pool", "pool_surface_m2, pool_volume_m3, filter_diameter, filter_count, motor_count", "Physical dimensions and hydraulic filtration capacity."],
        ["Categorical", "pool_type (community, outdoor, heated), deck_type (grass, paved, mixed)", "One-hot encoded categorical indicators."],
        ["Autoregressive Lags", "ph_lag1, ph_lag2, chlorine_lag1, chlorine_lag2, turbidity_lag1, turbidity_lag2", "Chemical measurements at the preceding 1 and 2 visits."],
        ["Rolling Stats", "ph_roll3_mean, ph_roll3_std, chlorine_roll3_mean, chlorine_roll3_std, turbidity_roll3_mean", "3-visit rolling mean and standard deviation capturing baseline stability."],
        ["Temporal & Calendar", "days_since_last_visit, visit_month, visit_is_summer, visit_day_of_week, visit_year, pool_visit_number", "Inter-visit gap k, summer indicator (June–Sept), and bather surge cycle."],
        ["Equipment Controls", "hypochlorite_dosing_pct, hypochlorite_dosing_hours, ph_dosing_pct, ph_dosing_hours, daily_filtration_hours, water_temperature", "Pump output percentages and daily filtration run times."],
        ["Product History", "last_total_chlorine_applied, total_ph_minus_product", "Total kg of chlorine and acid products applied at previous visit."],
        ["Regulatory Headroom", "chlorine_headroom_low (Cl - 0.5), chlorine_headroom_high (5.0 - Cl), ph_headroom_low (pH - 7.2), ph_headroom_high (8.0 - pH), turbidity_headroom (5.0 - Turb), min_headroom", "Linear distance to nearest regulatory boundary under RD 742/2013."],
        ["Client Target Distances", "cl_below_client_target (max(0, 1.0 - Cl)), cl_above_client_target (max(0, Cl - 1.5))", "Linear shortfall or surplus relative to client 1.0–1.5 mg/L target."],
        ["Trends & Drift Rates", "ph_trend, chlorine_trend, turbidity_trend, ph_rate_per_day, chlorine_rate_per_day, turbidity_rate_per_day", "Directional delta and velocity of chemical change per elapsed day."],
        ["Breach Compliance", "consecutive_clean_visits, breach_rate_last5, current_any_breach, current_ph_breach, current_chlorine_breach, multi_visit_day", "Historical reliability metrics and chronic problem-pool indicators."],
        ["Water Chemistry", "ph_deviation (|pH - 7.4|), chlorine_deficit (max(0, 0.5 - Cl)), cl_effectiveness_index (HOCl active fraction), chlorine_dose_per_m3, ph_minus_dose_per_m3, chlorine_decay_per_m3", "Physical chemical balance and volume-normalized dosing rates."],
        ["Setpoint Degradation", "setpoint_free_chlorine (2.5), setpoint_ph (7.4), setpoint_turbidity (0.5), cl_degradation_from_setpoint, ph_drift_from_setpoint, turb_accumulation_from_setpoint, and respective per-day rates", "Quantifies the treat -> degrade -> re-measure cycle from post-treatment ideal."],
        ["Weather Intelligence", "9 current weather + 4 cumulative weather + 9 tomorrow forecast features", "Atmospheric photolysis, ambient temperature, rainfall, and wind signals."],
    ]
    add_styled_table(doc, feat_headers, feat_data, [Inches(1.5), Inches(2.2), Inches(2.8)])

    # -----------------------------------------------------------------------
    # CHAPTER 5: MACHINE LEARNING MODELS
    # -----------------------------------------------------------------------
    add_heading_1(doc, "5. Machine Learning Models, Training & Evaluation")

    add_heading_2(doc, "5.1 Next-Day Synthetic Target Formulation")
    add_paragraph(doc, "Because technician visits are spaced k days apart (median k = 3 days), raw next-visit readings cannot be used directly for daily dispatching. The target represents the chemical state on the NEXT CALENDAR DAY (T + 1), linearly interpolated from the assumed post-treatment setpoint:")
    
    add_code_block(doc, "Target_Tomorrow = Setpoint + (Reading_Next_Visit - Setpoint) * (1 / k)")
    
    add_bullet(doc, "When k = 1 (consecutive day visit), Target_Tomorrow is the exact measured next reading.", "Consecutive Visits (k=1): ")
    add_bullet(doc, "When k = 3, Target_Tomorrow represents exactly 1 day of degradation from the post-treatment setpoint.", "Multi-Day Gaps (k>1): ")
    add_bullet(doc, "When no subsequent visit exists (end of time-series), Target_Tomorrow defaults to 1 day of physical kinetic decay.", "Terminal Visits (NaN gap): ")

    add_heading_2(doc, "5.2 Chronological 80/20 Train/Test Split")
    add_paragraph(doc, "To strictly avoid temporal data leakage, an 80/20 chronological split is enforced at the 80th percentile timestamp (October 13, 2025):")
    add_bullet(doc, "28,910 rows covering January 2, 2023 through October 13, 2025.", "Training Partition (80%): ")
    add_bullet(doc, "7,228 rows covering October 14, 2025 through August 5, 2026.", "Test Holdout Partition (20%): ")

    add_heading_2(doc, "5.3 XGBoost Model Configurations & Loss Function")
    add_paragraph(doc, "Three dedicated XGBoost regressors (Model A: Free Chlorine, Model C: pH, Model D: Turbidity) are trained with shared hyper-parameters:")
    add_bullet(doc, "n_estimators=500, max_depth=5, learning_rate=0.05, subsample=0.8, colsample_bytree=0.8, reg_alpha=0.1, reg_lambda=1.0, random_state=42.", "Tree Hyper-Parameters: ")
    add_bullet(doc, "Early stopping on holdout RMSE with a patience of 50 rounds.", "Early Stopping: ")
    add_bullet(doc, "Sample weights for Free Chlorine are amplified by 3.0x on rows where a regulatory breach occurred, forcing the trees to prioritize extreme hazard conditions.", "Breach Sample Weighting: ")

    add_heading_2(doc, "5.4 Verified Model Performance Highlights")
    perf_headers = ["Model", "Target Variable", "MAE", "RMSE", "R² Score", "P90 Error", "Best Iteration"]
    perf_data = [
        ["Model A", "Free Chlorine Tomorrow (mg/L)", "0.1972", "0.3447", "0.2571", "0.4503 mg/L", "118 trees"],
        ["Model C", "pH Tomorrow (pH units)", "0.0332", "0.0538", "0.2974", "0.0811 pH", "52 trees"],
        ["Model D", "Turbidity Tomorrow (NTU)", "0.0420", "0.0777", "0.4013", "0.0940 NTU", "21 trees"],
    ]
    add_styled_table(doc, perf_headers, perf_data, [Inches(1.1), Inches(2.2), Inches(0.8), Inches(0.8), Inches(0.8), Inches(1.1), Inches(0.9)])

    add_paragraph(doc, "Operational Precision Context: A pH MAE of 0.0332 is well within standard handheld electronic pH meter accuracy (±0.10 pH units). Free Chlorine MAE of 0.1972 mg/L provides tight tracking for dispatching dosing adjustments before pathogen (0.5 mg/L) or closure (5.0 mg/L) thresholds are breached.")

    # -----------------------------------------------------------------------
    # CHAPTER 6: PHYSICAL KINETICS RATE INTEGRATION
    # -----------------------------------------------------------------------
    add_heading_1(doc, "6. Physical Kinetics Rate Integration Engine")

    add_paragraph(doc, "A pure machine learning model trained on historical data can suffer from non-physical extrapolations under extreme weather. The V6 system implements a hybrid physics-ML integration engine that couples XGBoost predictions with first-principles physical chemistry:")

    add_heading_2(doc, "6.1 UV Photolysis First-Order Chlorine Decay")
    add_paragraph(doc, "In the absence of newly applied chlorine, hypochlorous acid (HOCl) decomposes under ultraviolet light according to first-order solar decay kinetics:")
    add_code_block(doc, 
        "decay_k = 0.15 + 0.003 * max(0, solar_radiation - 15.0)\n"
        "Cl_kinetic = Anchor_Cl * exp(-decay_k / 3.0)\n"
        "Pred_Cl = max(0.0, min(Raw_XGB_Cl, Cl_kinetic))"
    )
    add_paragraph(doc, "Mechanism: Solar radiation above 15 MJ/m² linearly accelerates decay. If no chemical product was added, the predicted chlorine is bounded by the physical kinetic ceiling.")

    add_heading_2(doc, "6.2 Temperature-Driven CO2 Degassing & Carbonate pH Drift")
    add_paragraph(doc, "Water exposed to ambient aeration and elevated summer temperature loses dissolved carbon dioxide, causing a natural upward drift in pH:")
    add_code_block(doc, 
        "daily_ph_drift = 0.035 + 0.0015 * max(0, temp_max - 25.0)\n"
        "pH_kinetic = Anchor_pH + daily_ph_drift\n"
        "Pred_pH = min(8.6, max(Raw_XGB_pH, pH_kinetic))"
    )
    add_paragraph(doc, "Mechanism: In Spanish summer heat (>25°C), pH naturally drifts upward at +0.035 to +0.060 units/day. Unless acid (pH Minus) is applied, the prediction is bounded below by this thermodynamic baseline.")

    add_heading_2(doc, "6.3 Wind-Borne Environmental Turbidity Accumulation")
    add_paragraph(doc, "Atmospheric dust and wind-borne debris continuously deposit particulate matter into outdoor pools:")
    add_code_block(doc, 
        "daily_turb_rise = 0.045 + 0.002 * max(0, wind_speed_max - 10.0)\n"
        "Turb_kinetic = Anchor_Turb + daily_turb_rise\n"
        "Pred_Turb = min(5.0, max(Raw_XGB_Turb, Turb_kinetic))"
    )

    # -----------------------------------------------------------------------
    # CHAPTER 7: CHAINED MULTI-STEP FORECASTING
    # -----------------------------------------------------------------------
    add_heading_1(doc, "7. Chained Multi-Step Daily Forecasting Engine")

    add_paragraph(doc, "When an operator opens the dashboard on Thursday for a pool last visited on Monday, a single 1-step prediction is useless. The Chained Multi-Step Predictor ('ml/inference/predictor.py') executes a sequential rollout across all intervening calendar days:")

    add_code_block(doc,
        "Monday (Actual Last Visit Measured Reading)\n"
        "  └── Step 1: Predict Tuesday   (Inject Tuesday actual weather; update lags & setpoint drift)\n"
        "        └── Step 2: Predict Wednesday (Inject Wednesday actual weather; roll state forward)\n"
        "              └── Step 3: Predict Thursday [TODAY] (Inject Thursday live weather)\n"
        "                    └── Step 4: Predict Friday [TOMORROW] (Inject Friday Open-Meteo forecast)"
    )

    add_heading_2(doc, "7.1 Dynamic State Recalculation at Each Step")
    add_paragraph(doc, "At each iteration t -> t + 1:")
    add_bullet(doc, "The predicted chemical concentrations (Cl_t, pH_t, Turb_t) become the inputs for step t + 1.", "Autoregressive Feedback: ")
    add_bullet(doc, "chlorine_lag1, ph_lag1, rolling 3-visit means, and standard deviations are dynamically updated.", "Lag & Rolling Updates: ")
    add_bullet(doc, "Headrooms (chlorine_headroom_low, etc.) and trend rates are recalculated against the new state.", "Headroom & Trend Recalculation: ")
    add_bullet(doc, "The exact historical or forecasted weather for day t and t + 1 is retrieved from the weather cache.", "Weather Injection: ")

    add_heading_2(doc, "7.2 Uncertainty Fan-Out Estimation")
    add_paragraph(doc, "Because forecast uncertainty compounds over time, each day step is assigned an UncertaintyBand:")
    add_code_block(doc, "Band_Width(step) = Base_MAE * sqrt(step)")
    add_paragraph(doc, "This provides operators with high-confidence predictions for Today/Tomorrow and transparent error expansion for extended 7-day horizons.")

    add_heading_2(doc, "7.3 Operational Urgency Classification")
    urg_headers = ["Classification Tier", "Trigger Condition", "Operational Action Required"]
    urg_data = [
        ["🚨 Immediate / URGENT", "Predicted Cl < 0.5 mg/L or > 5.0 mg/L\nPredicted pH < 7.2 or > 8.0\nPredicted Turbidity > 5.0 NTU", "Immediate technician dispatch required today. Severe pathogen or chemical hazard."],
        ["⚠️ Advised (Maintenance)", "Predicted Cl < 1.0 mg/L (Breaches client target minimum)", "Schedule routine visit within 24–48 hours to prevent regulatory breach."],
        ["⚠️ Monitor (Overdose Buffer)", "Predicted Cl > 2.0 mg/L (within Spanish Mediterranean buffer)", "No immediate hazard; recommend reducing pump dosing % at next visit."],
        ["✅ Routine (Optimal)", "Cl in 1.0–1.5 mg/L, pH in 7.2–8.0, Turbidity ≤ 1.0 NTU", "Pool operates in optimal equilibrium. No technician intervention needed."],
    ]
    add_styled_table(doc, urg_headers, urg_data, [Inches(1.8), Inches(2.7), Inches(2.0)])

    # -----------------------------------------------------------------------
    # CHAPTER 8: CHEMICAL DOSING OPTIMIZER
    # -----------------------------------------------------------------------
    add_heading_1(doc, "8. Automated Chemical Dosing Optimization Engine")

    add_paragraph(doc, "The Dosing Optimizer ('ml/inference/optimiser.py') performs an exhaustive grid search over dosing pump configurations to find the minimal chemical effort required to keep water parameters within optimal client targets [Cl 1.0–1.5 mg/L, pH 7.2–8.0]:")

    add_heading_2(doc, "8.1 Grid Search Formulation (525 Candidate Configurations)")
    add_bullet(doc, "0% to 100% in 5% increments (21 discrete levels).", "Hypochlorite Dosing Percentage: ")
    add_bullet(doc, "0.0 to 24.0 hours in 1.0-hour increments (25 discrete levels).", "Pump Operating Duration: ")
    add_bullet(doc, "21 × 25 = 525 candidate configurations evaluated per pool in <15 milliseconds.", "Total Candidate Space: ")

    add_heading_2(doc, "8.2 Multi-Objective Loss Function")
    add_paragraph(doc, "For each candidate (Dosing%, Hours), the model predicts next-day chemical outcomes and evaluates:")
    add_code_block(doc,
        "cl_penalty = max(0, 1.0 - pred_cl) + max(0, pred_cl - 1.5)\n"
        "ph_penalty = max(0, 7.2 - pred_ph) + max(0, pred_ph - 8.0)\n"
        "total_penalty = cl_penalty + ph_penalty\n"
        "dosing_cost = (hypochlorite_dosing_pct / 100.0) * hypochlorite_dosing_hours"
    )
    add_paragraph(doc, "Grid points are sorted primarily by ascending 'total_penalty' (0.0 = perfect compliance) and secondarily by ascending 'dosing_cost' (minimal chemical and pump wear).")

    # -----------------------------------------------------------------------
    # CHAPTER 9: SHAP EXPLAINABILITY
    # -----------------------------------------------------------------------
    add_heading_1(doc, "9. SHAP Explainability & Top Global Feature Drivers")

    add_paragraph(doc, "TreeExplainer Shapley value analysis was conducted on all three models across the holdout test set to ensure total transparency and interpretability:")

    shap_headers = ["Target Model", "Top 5 SHAP Feature Drivers", "Mean |SHAP| Value", "Physical & Chemical Rationale"]
    shap_data = [
        ["Free Chlorine (Model A)", "1. chlorine_headroom_low\n2. chlorine_roll3_mean\n3. visit_is_summer\n4. cl_degradation_rate_from_setpoint\n5. chlorine_headroom_high", "0.0481\n0.0469\n0.0293\n0.0236\n0.0152", "Proximity to minimum 0.5 mg/L threshold, baseline rolling stability, summer photolysis surge, and setpoint decay rate dominate chlorine consumption."],
        ["pH (Model C)", "1. ph_roll3_mean\n2. ph_headroom_low\n3. ph_drift_rate_from_setpoint\n4. visit_is_summer\n5. visit_year", "0.0154\n0.0096\n0.0068\n0.0039\n0.0027", "Rolling pH baseline, distance from 7.2 acidic boundary, and degassing rate from setpoint govern carbonate balance."],
        ["Turbidity (Model D)", "1. turb_accumulation_rate_from_setpoint\n2. turbidity_roll3_mean\n3. visit_is_summer\n4. visit_day_of_week\n5. w_tmrw_temp_mean", "0.0123\n0.0116\n0.0115\n0.0058\n0.0033", "Accumulation rate from 0.5 NTU setpoint is the #1 driver, followed by summer dust and weekend bather surges."],
    ]
    add_styled_table(doc, shap_headers, shap_data, [Inches(1.5), Inches(2.2), Inches(1.1), Inches(2.0)])

    add_heading_2(doc, "9.1 SHAP Summary Charts")
    add_image_with_caption(doc, PROJECT_ROOT / "outputs" / "shap_summary_chlorine_next.png", "Figure 1: SHAP Feature Importance Summary — Next-Day Free Chlorine Model")
    add_image_with_caption(doc, PROJECT_ROOT / "outputs" / "shap_summary_ph_next.png", "Figure 2: SHAP Feature Importance Summary — Next-Day pH Model")
    add_image_with_caption(doc, PROJECT_ROOT / "outputs" / "shap_summary_turbidity_next.png", "Figure 3: SHAP Feature Importance Summary — Next-Day Turbidity Model")

    # -----------------------------------------------------------------------
    # CHAPTER 10: FULL-STACK ARCHITECTURE & IMPLEMENTATION
    # -----------------------------------------------------------------------
    add_heading_1(doc, "10. Full-Stack System Architecture & Production Implementation")

    add_paragraph(doc, "The system is implemented as an enterprise-grade, asynchronous full-stack platform designed for Dockerized cloud or on-premises deployment:")

    add_heading_2(doc, "10.1 System Topology & Component Interactions")
    add_bullet(doc, "High-performance Python 3.10+ ASGI service using Pydantic V2 validation, CORS middleware, and request tracing.", "Backend API (FastAPI): ")
    add_bullet(doc, "Type-safe asynchronous Prisma Client Python connected to PostgreSQL 16.", "Database ORM (Prisma): ")
    add_bullet(doc, "In-process AsyncIOScheduler managing daily weather synchronization (4:00 AM) and periodic model retraining (Monday 3:00 AM).", "Job Scheduler (APScheduler): ")
    add_bullet(doc, "Single-page operator application built with React 19, TypeScript, Vite, TanStack React Query, Recharts, and Tailwind CSS.", "Frontend Dashboard (React 19): ")

    add_heading_2(doc, "10.2 PostgreSQL Relational Database Schema")
    add_paragraph(doc, "The database schema ('prisma/schema.prisma') defines five core relational models:")
    add_bullet(doc, "pool_id (PK), community_name, pool_type, deck_type, pool_volume_m3, pool_surface_m2, filter_diameter, filter_count, motor_count, pool flags (heated, outdoor, skimmer, overflow, rectangular).", "Pool (pools table): ")
    add_bullet(doc, "id (PK autoincrement), pool_id (FK cascade), reading_date, technician, ph, free_chlorine, turbidity, hypochlorite_dosing_pct, hypochlorite_dosing_hours, water_temperature, source, created_at. Unique compound constraint on [pool_id, reading_date].", "Reading (readings table): ")
    add_bullet(doc, "date (PK DateTime), w_temp_max, w_temp_mean, w_uv_max, w_solar_radiation, w_sunshine_hours, w_precipitation_mm, w_wind_max_kmh, w_et0, w_weather_code, fetched_at.", "WeatherDaily (weather_daily table): ")
    add_bullet(doc, "run_id (PK String), artifact_dir, created_at, is_active (Int), metrics_json, feature_schema_json, promoted_at, promote_reason.", "ModelRun (model_runs table): ")
    add_bullet(doc, "id (PK autoincrement), source, filename, pool_count, row_count, skipped_count, created_at, detail_json.", "IngestLog (ingest_logs table): ")

    add_heading_2(doc, "10.3 Background Job Scheduler & Automated Promotion Gate")
    add_paragraph(doc, "1. Daily Weather Refresh: Runs every morning at 4:00 AM ('0 4 * * *'). Fetches yesterday's historical weather plus 7-day forward forecasts from Open-Meteo, upserts into 'weather_daily', and invalidates the in-memory cache.")
    add_paragraph(doc, "2. Weekly Model Retraining & Promotion: Runs weekly on Monday at 3:00 AM ('0 3 * * 1'). Spawns 'ml.training.train' as a non-blocking subprocess if at least 200 new reading rows have been ingested since the last active model run.")
    add_paragraph(doc, "Automated Promotion Gate: The newly trained candidate model is promoted to active status only if its holdout MAE is no worse than the active run within strict tolerances (Cl MAE slack: 0.02 mg/L, pH MAE slack: 0.005, Turbidity MAE slack: 0.01 NTU). If promoted, 'models/latest.json' and PostgreSQL 'is_active' are updated atomically, and the PredictionService hot-reloads without server downtime.")

    add_heading_2(doc, "10.4 Multi-Channel Data Ingestion")
    add_bullet(doc, "Web UI drag-and-drop supporting .xlsx and .csv files with automated fuzzy column detection, interactive header mapping, and instant preview.", "Interactive File Upload: ")
    add_bullet(doc, "Clean web interface for technicians to log daily field measurements on mobile devices or tablets.", "Manual Reading Form: ")
    add_bullet(doc, "Secure endpoint ('POST /api/ingest/readings') supporting bulk JSON payloads with optional Bearer token authentication.", "Authenticated REST JSON Ingest: ")
    add_bullet(doc, "Direct multipart upload endpoint ('POST /api/ingest/readings/file') for automated script or cron feeds.", "Automated File Ingest: ")

    # -----------------------------------------------------------------------
    # CHAPTER 11: REST API SPECIFICATION
    # -----------------------------------------------------------------------
    add_heading_1(doc, "11. Complete REST API Reference")

    api_headers = ["Endpoint Route", "Method", "Auth", "Description & Response Schema"]
    api_data = [
        ["/healthz / /healthz/live", "GET", "None", "Kubernetes liveness probe. Returns {status: 'ok', timestamp: ISO8601}."],
        ["/healthz/ready", "GET", "None", "Readiness probe verifying DB connection, active model status, and weather cache freshness."],
        ["/api/fleet", "GET", "None", "Paginated fleet overview with query date filtering, search (q), urgency filtering, and Today/Tomorrow forecast chips."],
        ["/api/fleet/pool-ids", "GET", "None", "Returns an array of all active registered pool IDs."],
        ["/api/fleet/dates", "GET", "None", "Returns min date, max date, and total active dates in the fleet readings history."],
        ["/api/pool/{pool_id}", "GET", "None", "Complete pool analytics: latest measurements, chained multi-day forecast with uncertainty bands, history time-series, and dosing recommendation."],
        ["/api/optimise/{pool_id}", "GET", "None", "Executes 525-grid search for pool; returns recommended pump % and hours, predicted outcome, and top-3 configurations."],
        ["/api/upload", "POST", "None", "Uploads Excel/CSV file (max 15MB); returns auto-detected column mappings and 5-row preview."],
        ["/api/map-columns", "POST", "None", "Confirms column mapping for upload_id; imports valid rows into PostgreSQL and records ingest log."],
        ["/api/readings", "POST", "None", "Accepts single manual technician reading; validates chemical ranges and upserts into database."],
        ["/api/ingest/readings", "POST", "Bearer (opt)", "Programmatic ingestion endpoint accepting JSON array of reading objects."],
        ["/api/ingest/readings/file", "POST", "Bearer (opt)", "Programmatic multipart file ingest endpoint with auto-mapping."],
        ["/api/admin/runs", "GET", "Admin Token", "Lists historical model runs, test metrics, promotion timestamps, and active status."],
        ["/api/admin/retrain", "POST", "Admin Token", "Triggers manual retraining pipeline subprocess and hot-reloads prediction service."],
        ["/api/admin/weather-status", "GET", "Admin Token", "Returns latest weather synchronization timestamp."],
        ["/api/admin/weather-refresh", "POST", "Admin Token", "Triggers immediate weather synchronization from Open-Meteo."],
        ["/api/admin/ingest-log", "GET", "Admin Token", "Returns audit log of all file uploads, API ingests, and manual submissions."],
    ]
    add_styled_table(doc, api_headers, api_data, [Inches(1.8), Inches(0.8), Inches(1.0), Inches(2.9)])

    # -----------------------------------------------------------------------
    # CHAPTER 12: DEPLOYMENT & TESTING
    # -----------------------------------------------------------------------
    add_heading_1(doc, "12. Deployment, Containerization & Verification")

    add_heading_2(doc, "12.1 Docker Multi-Container Architecture")
    add_paragraph(doc, "The repository provides a complete 'docker-compose.yml' orchestrating four interconnected container services:")
    add_bullet(doc, "PostgreSQL 16 Alpine container with persistent named volume 'postgres_data' and internal healthchecks.", "1. Database (postgres): ")
    add_bullet(doc, "Runs 'python -m backend.store.migrate' on startup to push schema definitions and seed initial pool records.", "2. Migration Seeder (migrate): ")
    add_bullet(doc, "FastAPI server running on port 8000 with volume mounts for live model artifacts ('models/'), data caches, and outputs.", "3. Backend API (backend): ")
    add_bullet(doc, "Nginx container serving pre-compiled React 19 production assets on port 8080 with reverse-proxy routing to the API backend.", "4. Frontend Web (frontend): ")

    add_heading_2(doc, "12.2 Automated Verification & Test Suite")
    add_paragraph(doc, "The codebase includes a comprehensive Pytest test suite ('tests/') validating all critical subsystems:")
    add_bullet(doc, "Validates that feature vectors generated by 'ml/features.py' match 'ml/training/steps.py' exactly with zero drift.", "Feature Parity Tests: ")
    add_bullet(doc, "Tests 'predict_forward' against deterministic mock weather inputs to verify boundary conditions and uncertainty bands.", "Inference Chaining Tests: ")
    add_bullet(doc, "Validates that candidate models are accepted or rejected strictly according to configured MAE slacks.", "Promotion Gate Tests: ")
    add_bullet(doc, "Tests FastAPI endpoints, parameter validation, error responses, and Bearer token security.", "API Endpoint Tests: ")
    add_bullet(doc, "Validates Prisma schema synchronization, batch upserts, and cascading delete integrity.", "Store & Database Tests: ")

    # -----------------------------------------------------------------------
    # CHAPTER 13: CONCLUSION & SUMMARY
    # -----------------------------------------------------------------------
    add_heading_1(doc, "13. Conclusion & Strategic Operational Summary")

    add_paragraph(doc, "The Spain (Alicante) Collective-Use Swimming Pool Predictive Maintenance System (V6.0) represents a complete transition from reactive, heuristic pool management to an automated, scientifically grounded predictive operation:")
    add_bullet(doc, "Bridges the gap between physical chemical kinetics and machine learning, eliminating non-physical predictions.", "1. Physical-ML Hybridization: ")
    add_bullet(doc, "Aligns targets with real-world field workflows, reducing Free Chlorine MAE to 0.1972 mg/L.", "2. Post-Treatment Setpoint Re-Anchoring: ")
    add_bullet(doc, "Eliminates inter-visit blindspots by projecting pool water quality day-by-day up to tomorrow and beyond.", "3. Multi-Day Chained Rollout: ")
    add_bullet(doc, "Evaluates 525 candidate pump configurations per pool to achieve maximum safety with minimal chemical spend.", "4. Automated Dosing Optimization: ")
    add_bullet(doc, "Equips dispatchers and technicians with intuitive Today/Tomorrow alerts, historical charts, and mobile logging.", "5. Production Full-Stack Readiness: ")

    # Save document
    doc.save(str(DOCX_OUTPUT_PATH))
    print(f"Documentation successfully generated at: {DOCX_OUTPUT_PATH}")


if __name__ == "__main__":
    build_complete_document()

